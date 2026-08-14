#!/usr/bin/env python3
"""Build the comparative technical review as DOCX and Markdown.

The DOCX is the authoring source for the final PDF. The script uses only local,
user-supplied sources plus the packaged bibliography.
"""
from __future__ import annotations

import csv
import json
import math
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "00_Report"
DATA = ROOT / "02_Data"
FIG = ROOT / "03_Figures"
REFS = ROOT / "06_References"
REPORT.mkdir(exist_ok=True)

DOCX_PATH = REPORT / "Klootwijk_Generational_Topology_Comparative_Review.docx"
MD_PATH = REPORT / "Klootwijk_Generational_Topology_Comparative_Review.md"

# Theme
NAVY = "0B2A3D"
TEAL = "148C9B"
GOLD = "D7A51F"
CORAL = "C8665A"
GREEN = "3B8D70"
PURPLE = "695C9C"
LIGHT = "EAF2F4"
LIGHT_GOLD = "FFF7DD"
LIGHT_CORAL = "FBEDEA"
MID = "B6C8CF"
DARK = "243640"
GREY = "6E7C84"
WHITE = "FFFFFF"


def rgb(hexstr: str) -> RGBColor:
    return RGBColor.from_string(hexstr)


def set_cell_shading(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_row_cant_split(row):
    """Keep a table row intact across pages for readable continuation."""
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    cant.set(qn("w:val"), "true")
    trPr.append(cant)


def set_cell_width(cell, width_cm: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_table_borders(table, color=MID, size=5):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": size, "color": color},
                bottom={"val": "single", "sz": size, "color": color},
                left={"val": "single", "sz": size, "color": color},
                right={"val": "single", "sz": size, "color": color},
            )


def set_table_no_outer_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"}
            )


def add_run_with_citations(paragraph, text: str, bold=False, italic=False, color=None, size=None, font=None):
    """Add text; keep bracketed citation chunks non-breaking where practical."""
    # A simple splitter keeps citations visually coherent without complex hyperlinking.
    parts = []
    cursor = 0
    while cursor < len(text):
        start = text.find("[", cursor)
        if start == -1:
            parts.append((text[cursor:], False)); break
        end = text.find("]", start)
        if end == -1:
            parts.append((text[cursor:], False)); break
        if start > cursor:
            parts.append((text[cursor:start], False))
        parts.append((text[start:end+1], True))
        cursor = end + 1
    if not parts:
        parts = [(text, False)]
    for part, is_cite in parts:
        r = paragraph.add_run(part)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = rgb(color)
        if size:
            r.font.size = Pt(size)
        if font:
            r.font.name = font
        if is_cite:
            r.font.color.rgb = rgb(TEAL)
            r.font.size = Pt((size or 10.2) - 0.3)


class Builder:
    def __init__(self):
        self.doc = Document()
        self.md: list[str] = []
        self._setup()

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(1.75)
        sec.bottom_margin = Cm(1.65)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)
        sec.header_distance = Cm(0.65)
        sec.footer_distance = Cm(0.65)
        sec.different_first_page_header_footer = True

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Inter"
        normal.font.size = Pt(9.6)
        normal.font.color.rgb = rgb(DARK)
        normal.paragraph_format.space_after = Pt(5.2)
        normal.paragraph_format.line_spacing = 1.12
        normal.paragraph_format.widow_control = True

        for name, size, color, before, after in [
            ("Title", 28, NAVY, 0, 10),
            ("Subtitle", 14.5, TEAL, 0, 8),
            ("Heading 1", 20, NAVY, 18, 8),
            ("Heading 2", 14.5, TEAL, 13, 5),
            ("Heading 3", 11.5, CORAL, 9, 3),
        ]:
            st = styles[name]
            st.font.name = "Inter"
            st.font.size = Pt(size)
            st.font.bold = True if "Heading" in name or name == "Title" else False
            st.font.color.rgb = rgb(color)
            st.paragraph_format.space_before = Pt(before)
            st.paragraph_format.space_after = Pt(after)
            st.paragraph_format.keep_with_next = True

        cap = styles["Caption"]
        cap.font.name = "Inter"
        cap.font.size = Pt(8.1)
        cap.font.italic = True
        cap.font.color.rgb = rgb(GREY)
        cap.paragraph_format.space_before = Pt(3)
        cap.paragraph_format.space_after = Pt(7)
        cap.paragraph_format.keep_with_next = False

        # Custom styles
        def make_style(name, base="Normal"):
            try:
                return styles[name]
            except KeyError:
                return styles.add_style(name, 1)

        lead = make_style("Lead")
        lead.font.name = "Noto Serif"
        lead.font.size = Pt(11.7)
        lead.font.color.rgb = rgb(NAVY)
        lead.paragraph_format.line_spacing = 1.18
        lead.paragraph_format.space_after = Pt(8)

        small = make_style("Small")
        small.font.name = "Inter"
        small.font.size = Pt(8)
        small.font.color.rgb = rgb(GREY)
        small.paragraph_format.space_after = Pt(3)
        small.paragraph_format.line_spacing = 1.02

        code = make_style("Code")
        code.font.name = "DejaVu Sans Mono"
        code.font.size = Pt(8.7)
        code.font.color.rgb = rgb(NAVY)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.left_indent = Cm(0.4)

        quote = make_style("Quote")
        quote.font.name = "Noto Serif"
        quote.font.size = Pt(10.8)
        quote.font.italic = True
        quote.font.color.rgb = rgb(NAVY)
        quote.paragraph_format.left_indent = Cm(0.55)
        quote.paragraph_format.right_indent = Cm(0.35)
        quote.paragraph_format.space_before = Pt(7)
        quote.paragraph_format.space_after = Pt(7)

        # Header/footer
        for section in self.doc.sections:
            hp = section.header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = hp.add_run("KLOOTWIJK GENERATIONAL TOPOLOGY REVIEW")
            r.font.name = "Inter"; r.font.size = Pt(7.5); r.font.bold = True; r.font.color.rgb = rgb(TEAL)
            fp = section.footer.paragraphs[0]
            add_page_number(fp)
            for run in fp.runs:
                run.font.name = "Inter"; run.font.size = Pt(8); run.font.color.rgb = rgb(GREY)

    def para(self, text: str, style=None, bold_prefix: str | None = None, align=None):
        p = self.doc.add_paragraph(style=style)
        if align is not None:
            p.alignment = align
        if bold_prefix and text.startswith(bold_prefix):
            r = p.add_run(bold_prefix)
            r.bold = True
            r.font.color.rgb = rgb(NAVY)
            add_run_with_citations(p, text[len(bold_prefix):])
        else:
            add_run_with_citations(p, text)
        self.md.append(text + "\n")
        return p

    def heading(self, text: str, level=1, page_break=False):
        if page_break:
            self.doc.add_page_break()
        p = self.doc.add_heading(text, level=level)
        # Accent rule under H1
        if level == 1:
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "14")
            bottom.set(qn("w:space"), "5")
            bottom.set(qn("w:color"), TEAL)
            pBdr.append(bottom)
            pPr.append(pBdr)
        self.md.append("#"*level + " " + text + "\n")
        return p

    def bullets(self, items: Sequence[str], level=0, numbered=False):
        for i, item in enumerate(items, start=1):
            style = "List Number" if numbered else "List Bullet"
            p = self.doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Cm(0.65 + level*0.4)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            p.paragraph_format.space_after = Pt(2.7)
            add_run_with_citations(p, item)
            prefix = f"{i}. " if numbered else "- "
            self.md.append(prefix + item + "\n")
        self.md.append("\n")

    def callout(self, title: str, body: str, color=TEAL, fill=LIGHT):
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        cell = t.cell(0,0)
        set_cell_shading(cell, fill)
        set_cell_border(cell,
                        top={"val":"single","sz":12,"color":color},
                        bottom={"val":"single","sz":12,"color":color},
                        left={"val":"single","sz":12,"color":color},
                        right={"val":"single","sz":12,"color":color})
        set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title.upper())
        r.bold=True; r.font.name="Inter"; r.font.size=Pt(9.2); r.font.color.rgb=rgb(color)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        add_run_with_citations(p2, body, size=10.0)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)
        self.md.append(f"> **{title}.** {body}\n\n")

    def equation(self, text: str, label: str | None = None):
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell=t.cell(0,0)
        set_cell_shading(cell, "F5F8F9")
        set_cell_border(cell,
                        top={"val":"single","sz":6,"color":MID},
                        bottom={"val":"single","sz":6,"color":MID},
                        left={"val":"single","sz":6,"color":MID},
                        right={"val":"single","sz":6,"color":MID})
        set_cell_margins(cell, top=110,start=140,bottom=110,end=140)
        p=cell.paragraphs[0]
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(text)
        r.font.name="DejaVu Sans Mono";r.font.size=Pt(10.5);r.font.color.rgb=rgb(NAVY)
        if label:
            p2=cell.add_paragraph(label)
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p2.style=self.doc.styles["Small"]
        self.doc.add_paragraph().paragraph_format.space_after=Pt(0)
        self.md.append(f"`{text}`\n\n")

    def figure(self, path: Path, caption: str, width_cm: float = 16.5):
        p=self.doc.add_paragraph()
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run()
        r.add_picture(str(path), width=Cm(width_cm))
        cp=self.doc.add_paragraph(caption, style="Caption")
        cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        self.md.append(f"![{caption}](../03_Figures/{path.name})\n\n")

    def source_image(self, path: Path, caption: str, width_cm=10.0):
        self.figure(path, caption, width_cm)

    def table(self, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None,
              font_size=8.1, header_fill=NAVY, first_col_bold=True, alternate=True, repeat_header=True):
        table=self.doc.add_table(rows=1, cols=len(headers))
        table.alignment=WD_TABLE_ALIGNMENT.CENTER
        table.autofit=False
        hdr=table.rows[0]
        if repeat_header:
            set_repeat_table_header(hdr)
        set_row_cant_split(hdr)
        for j,h in enumerate(headers):
            cell=hdr.cells[j]
            set_cell_shading(cell, header_fill)
            set_cell_margins(cell, top=90,start=90,bottom=90,end=90)
            if widths: set_cell_width(cell,widths[j])
            p=cell.paragraphs[0]
            p.paragraph_format.space_after=Pt(0)
            r=p.add_run(str(h));r.bold=True;r.font.name="Inter";r.font.size=Pt(font_size);r.font.color.rgb=rgb(WHITE)
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for i,row in enumerate(rows):
            new_row = table.add_row()
            set_row_cant_split(new_row)
            cells=new_row.cells
            for j,val in enumerate(row):
                cell=cells[j]
                if widths: set_cell_width(cell,widths[j])
                if alternate and i%2==1: set_cell_shading(cell,"F2F6F7")
                set_cell_margins(cell, top=75,start=80,bottom=75,end=80)
                p=cell.paragraphs[0]
                p.paragraph_format.space_after=Pt(0)
                add_run_with_citations(p,str(val),bold=(first_col_bold and j==0),size=font_size)
                cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_table_borders(table)
        self.doc.add_paragraph().paragraph_format.space_after=Pt(0)
        # Markdown
        self.md.append("|"+"|".join(headers)+"|\n")
        self.md.append("|"+"|".join(["---"]*len(headers))+"|\n")
        for row in rows:
            self.md.append("|"+"|".join(str(v).replace("|","/").replace("\n"," ") for v in row)+"|\n")
        self.md.append("\n")
        return table

    def cards(self, cards: Sequence[tuple[str,str,str]], cols=2):
        rows=math.ceil(len(cards)/cols)
        t=self.doc.add_table(rows=rows,cols=cols)
        t.alignment=WD_TABLE_ALIGNMENT.CENTER
        t.autofit=True
        for idx,(title,body,color) in enumerate(cards):
            cell=t.cell(idx//cols,idx%cols)
            set_cell_shading(cell,"F8FAFB")
            set_cell_border(cell,
                            top={"val":"single","sz":8,"color":color},bottom={"val":"single","sz":8,"color":color},
                            left={"val":"single","sz":8,"color":color},right={"val":"single","sz":8,"color":color})
            set_cell_margins(cell,top=120,start=130,bottom=120,end=130)
            p=cell.paragraphs[0]
            r=p.add_run(title);r.bold=True;r.font.name="Inter";r.font.size=Pt(10);r.font.color.rgb=rgb(color)
            p2=cell.add_paragraph();p2.paragraph_format.space_after=Pt(0);add_run_with_citations(p2,body,size=8.9)
            self.md.append(f"**{title}.** {body}\n\n")
        # blank cell shading
        for idx in range(len(cards),rows*cols):
            cell=t.cell(idx//cols,idx%cols)
            set_cell_border(cell,top={"val":"nil"},bottom={"val":"nil"},left={"val":"nil"},right={"val":"nil"})
        self.doc.add_paragraph().paragraph_format.space_after=Pt(0)

    def page_break(self):
        self.doc.add_page_break()
        self.md.append("\n---\n\n")

    def save(self):
        # Metadata
        props=self.doc.core_properties
        props.title="From Stoichiometric Polymer Topology to Event-Driven Interface Topology"
        props.subject="Comparative technical review of Arie Klootwijk's patent and Tom Klootwijk's spherical substrate corpus"
        props.author="OpenAI - prepared from Klootwijk family sources"
        props.keywords="polyhydroxyether sulfone, topology, event calculus, membranes, biointerfaces, digital twin"
        self.doc.save(DOCX_PATH)
        MD_PATH.write_text("".join(self.md),encoding="utf-8")
        print(DOCX_PATH)
        print(MD_PATH)


def load_csv(path: Path) -> list[dict[str,str]]:
    with path.open(encoding="utf-8") as f:
        return list(csv.DictReader(f))


b=Builder()

# ---------------------------------------------------------------------------
# Cover
# ---------------------------------------------------------------------------
# top color band
band=b.doc.add_table(rows=1,cols=1)
band.alignment=WD_TABLE_ALIGNMENT.CENTER
cell=band.cell(0,0)
set_cell_shading(cell,NAVY)
set_cell_border(cell,top={"val":"nil"},bottom={"val":"nil"},left={"val":"nil"},right={"val":"nil"})
set_cell_margins(cell,top=200,start=140,bottom=200,end=140)
p=cell.paragraphs[0]
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("GENERATIONS OF TOPOLOGY")
r.font.name="Inter";r.font.size=Pt(12);r.bold=True;r.font.color.rgb=rgb(WHITE)

p=b.doc.add_paragraph()
p.paragraph_format.space_before=Pt(38)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("From Stoichiometric Polymer Topology\nto Event-Driven Interface Topology")
r.font.name="Inter";r.font.size=Pt(27);r.bold=True;r.font.color.rgb=rgb(NAVY)

p=b.doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("A comparative technical review of Arie Klootwijk's poly(hydroxy ether sulfone) patent and Tom Klootwijk's spherical substrate corpus")
r.font.name="Noto Serif";r.font.size=Pt(13.5);r.italic=True;r.font.color.rgb=rgb(TEAL)

# dual cards
b.cards([
    ("ARIE KLOOTWIJK - MOLECULAR TOPOLOGY", "Near-equimolar step-growth polymerization, solvent-controlled accessibility, linear-chain architecture, reactive hydroxyl interfaces, and experimentally specified process closure.", CORAL),
    ("TOM KLOOTWIJK - OPERATIONAL TOPOLOGY", "Local spherical support, compatibility sectors, guard crossings, transitions, uncertainty, event logs, and lineage in a query-first state architecture.", TEAL),
],cols=2)

p=b.doc.add_paragraph()
p.paragraph_format.space_before=Pt(26)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Prepared for Tom Klootwijk\n14 August 2026 | Version 1.0")
r.font.name="Inter";r.font.size=Pt(10);r.font.color.rgb=rgb(GREY)

p=b.doc.add_paragraph()
p.paragraph_format.space_before=Pt(22)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("UNOFFICIAL TECHNICAL REVIEW AND RESEARCH DESIGN PACKAGE")
r.bold=True;r.font.name="Inter";r.font.size=Pt(8.5);r.font.color.rgb=rgb(CORAL)

# first page footer disclaimer
p=b.doc.add_paragraph()
p.paragraph_format.space_before=Pt(28)
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("This report is not a certified patent translation, legal opinion, medical-device assessment, safety certification, or proof of patentability. It distinguishes source facts, established external methods, comparative inference, and research hypotheses.")
r.font.name="Inter";r.font.size=Pt(7.7);r.font.color.rgb=rgb(GREY)

b.md.extend([
    "# From Stoichiometric Polymer Topology to Event-Driven Interface Topology\n",
    "*A comparative technical review of Arie Klootwijk's poly(hydroxy ether sulfone) patent and Tom Klootwijk's spherical substrate corpus*\n\n",
    "Prepared for Tom Klootwijk, 14 August 2026.\n\n",
])
b.page_break()

# ---------------------------------------------------------------------------
# Executive verdict
# ---------------------------------------------------------------------------
b.heading("Executive verdict",1)
b.para("The two bodies of work are not the same invention expressed in different jargon. They operate at different physical and evidential layers. Arie Klootwijk's patent directly controls molecular graph topology and phase accessibility in a real polymerization. Tom Klootwijk's mature corpus primarily controls operational topology: which local states are relevant, which sectors may couple, when a guard becomes an event, how a transition is routed, and how lineage is preserved. The overlap is nevertheless substantial at the level of fundamental technique.",style="Lead")

b.callout("Central finding", "The strongest shared grammar is support/accessibility -> compatibility/balance -> local event -> connectivity or state transition -> lineage/closure. In Arie's work it is physically instantiated by solution polymerization; in Tom's work it is an explicit information architecture. The research opportunity is to connect the layers without mistaking analogy for mechanism.", color=TEAL, fill=LIGHT)

b.cards([
    ("Where Arie is stronger", "Matter, quantities, process windows, impurity control, kinetics/phase failure modes, isolation, specimens, and historical measurements. His work closes the physical loop.", CORAL),
    ("Where Tom is stronger today", "Query semantics, compatibility as a first-class operator, event/transition logic, uncertainty, lineage, local-to-global sensing, digital twins, and cross-domain portability.", TEAL),
    ("Where the overlap is deepest", "Both prevent useful structure from being lost by premature closure: Arie prevents crystallization or chain termination before high conversion; Tom resists premature rasterization/materialization before a query is resolved.", GOLD),
    ("Best combined direction", "A chemically addressable membrane or biointerface - potentially using a hydroxyl-rich sulfone polyether or safer analogue - coupled to a bounded, event-sourced digital twin with measured guards, uncertainty, and material lineage.", GREEN),
],cols=2)

b.heading("What 'better present day' means",2)
b.para("Tom's work is better positioned where the deliverable is a modern computational layer around complex physical interfaces: sparse sensing, state estimation, compatibility gates, event detection, provenance, and decision support. It is not yet better than the patent as a materials technology because the corpus does not supply a validated composition, manufacturing process, or measured device. The credible present-day advantage is therefore architectural rather than chemical: Tom can make the material system more queryable, auditable, adaptive, and multi-scale.")

b.bullets([
    "For computer graphics and simulation, the strongest target is a restricted equation-world substrate that answers state-at-time and next-event queries for controlled relation families, with rasterization or tracing retained as optional projection tools [S3, pp. 5-12].",
    "For membranes, the strongest target is event-centric monitoring and control of fouling, breakthrough, cleaning, drift, and morphology, connected to material and calibration lineage [S6, pp. 3-10; R07].",
    "For biointerfaces, the strongest target is a functionalized surface whose adhesion, aggregation, detachment, or binding events are captured sparsely and interpreted with topology descriptors and uncertainty [R08-R09].",
    "For material innovation, the inherited platform is most credible as a reactive thermoplastic coating, tie layer, blend modifier, or functionalization scaffold - not as an intrinsically ionic or automatically biocompatible membrane [S1, pp. 3-8; R02].",
])

b.heading("Decisions at a glance",2)
b.table(
    ["Question","Review decision","Why"],
    [
        ["Was the SO2/equimolar language a hidden cipher?","No evidence of ciphering.","SO2 is a neutral diaryl-sulfone bridge; equimolarity is the nonlinear molecular-weight control of bifunctional step-growth polymerization."],
        ["Is there a genuine topological overlap?","Yes, at a structural level.","Both systems gate local transformations by accessibility and compatibility, then update connectivity and history."],
        ["Is Tom's work more advanced?","At the operational-information layer, yes; at the materials-evidence layer, no.","The corpus expresses modern event, lineage, uncertainty, and sensing concepts; the patent physically demonstrates a process."],
        ["Can the work enable membranes and biointerfaces?","Yes as a research program, not as an established device.","Arie's polymer offers reactive surface handles; Tom offers a control/representation architecture. Transport, safety, and biological validation remain mandatory."],
        ["What should be built first?","Equation World Zero and Membrane World Zero.","They test the core architecture headlessly and against conventional baselines before expensive biomedical integration."],
    ],
    widths=[4.0,4.0,8.0],font_size=8.5
)

b.page_break()

# Contents and source map
b.heading("Contents and reading map",1)
contents=[
    ("1", "Scope, sources, and evidence discipline"),
    ("2", "Arie Klootwijk's technical core"),
    ("3", "Tom Klootwijk's mature technical core"),
    ("4", "Four layers of topology"),
    ("5", "Detailed fundamental crosswalk"),
    ("6", "Topological information science and computer graphics"),
    ("7", "Membranes and biointerfaces"),
    ("8", "Cross-domain applications and readiness"),
    ("9", "Research, validation, and IP roadmap"),
    ("10", "Conclusions"),
    ("A-D", "Formal notation, comparison/application tables, evidence ledger, references"),
]
b.table(["Part","Purpose"],contents,widths=[2.0,14.0],font_size=9.1,header_fill=TEAL)

b.heading("Source codes used in the report",2)
source_rows=[
    ["S1","SE 301 717 B original Swedish patent and the supplied English translation","Primary chemistry, examples, claims, and historical measurements"],
    ["S2","Arie Klootwijk SE301717B Technical Review","Prior chemical, historical, application, and safety analysis"],
    ["S3","Chronological Synthesis of the Spherical Substrate Line","Mature corpus interpretation, formal core, feasibility, and prototype discipline"],
    ["S4","Hollowland - Double Vacuum","Exploratory vocabulary and strong early claims; useful mainly when translated through S3/S6"],
    ["S5","21BenBurgersStrikesBackTelNetNiet","Exploratory folding, overlap, torque, and shared-domain motifs"],
    ["S6","Spherical Throughput: Practical Waveguide Liquid-Substrate Lensing","Bounded hardware translation, B.C.E. guards, metrics, failure criteria, and Hollowlens-0"],
]
b.table(["Code","Document","Role"],source_rows,widths=[1.3,6.2,8.5],font_size=8.4)

b.heading("Evidence tiers",2)
b.cards([
    ("TIER A - SOURCE FACT", "Directly stated or measured in the supplied patent or corpus.", NAVY),
    ("TIER B - ESTABLISHED PRINCIPLE", "Supported by standard polymer science, topology, controls, sensing, or official biomedical guidance.", TEAL),
    ("TIER C - COMPARATIVE INFERENCE", "A structural mapping proposed in this review; useful but not a historical or physical identity claim.", GOLD),
    ("TIER D - RESEARCH HYPOTHESIS", "A plausible application that requires experimental falsification and a matched baseline.", CORAL),
    ("TIER X - REJECTED OVERCLAIM", "A universal claim contradicted by physical constraints or demoted by the mature corpus itself.", PURPLE),
],cols=2)

b.callout("Interpretive rule", "The report preserves the corpus's own late correction: spheres, cones, SDF zero, one-bit, double vacuum, and hourglass are treated as typed operators in a shared state space unless a source specifies literal physics. Conversely, chemical terms retain their literal physical meaning and are not converted into computational symbols without an explicit model [S3, pp. 1-2].", color=GOLD, fill=LIGHT_GOLD)

# ---------------------------------------------------------------------------
# 1 Scope
# ---------------------------------------------------------------------------
b.heading("1. Scope, sources, and evidence discipline",1)
b.para("The review asks a precise question: do the fundamental techniques in Arie Klootwijk's polymer patent and Tom Klootwijk's topological-information corpus overlap, and if so, at what level? It then extends the comparison to present-day computer graphics, topological information science, membranes, biointerfaces, and other cross-domain applications. The answer depends on keeping four meanings of 'topology' separate: molecular connectivity, morphology and phase, interface and transport connectivity, and operational information topology.")

b.heading("1.1 Why a layered comparison is necessary",2)
b.para("A polymer chemist may use topology to distinguish linear, branched, cyclic, crosslinked, and network structures. A membrane scientist may use it for connected pores, tortuosity, percolating domains, or interface networks. A computer graphics researcher may use it for level sets, Reeb graphs, meshes, non-orientable manifolds, or topological descriptors. An information scientist may use it for locality, compatibility, routing, provenance, and global consistency. These are related but not interchangeable.")
b.para("The key risk in the corpus is a type error: a visual or algebraic metaphor can slide into a physical claim without passing through a model, observable, uncertainty, and validation step. The key risk in reading the patent is the opposite: reducing a sophisticated process architecture to ordinary practical chemistry and missing how carefully it controls graph connectivity, accessibility, and termination.")

b.figure(FIG/"figure_08_translation_ladder.png","Figure 1. Translation ladder used throughout the review. Abstraction is productive only after the physical system, governing model, observables, guard, transition, and lineage have been typed.",width_cm=16.2)

b.heading("1.2 What the sources support",2)
b.para("The patent supports a clear materials claim: a nearly linear, high-molecular-weight thermoplastic formed by reacting difunctional phenolic and epoxide compounds, with at least one monomer containing a diaryl-sulfone bridge. It supports a process claim: choose a sufficiently polar, substantially nonreactive solvent that keeps both monomers and the growing product in solution; use an alkaline catalyst; maintain close functional equivalence; exclude water and control purity; and allow the reaction to reach a target intrinsic viscosity [S1, pp. 3-10].")
b.para("The mature corpus supports a clear computational claim: the core should not be judged as rasterization, ray marching, voxel storage, or a conventional frame loop. Those are optional projections. The substrate is instead a finite grammar of directly queryable relations, phases, compatibility sectors, events, transitions, invariants, lineage, and an irreducible log. Its feasible form is explicitly restricted: fixed or bounded relation families, direct state/event queries, and benchmarked prototypes rather than a universal replacement for all simulation and memory [S3, pp. 1, 5-12].")
b.para("Spherical Throughput adds a hardware-facing discipline. It translates 'matrix-in-glass' into a calibrated transfer matrix, 'double vacuum' into deliberately uncoupled modes, a pinch point into a measured threshold or cutoff, and one-bit parity into a narrow route/validity flag. It rejects zero latency, zero heat, zero memory, and literal topological miracles, while retaining support, compatibility, guard crossings, uncertainty, energy, calibration, and failure criteria [S6, pp. 2-12].")

b.heading("1.3 What the sources do not support",2)
b.bullets([
    "No source supports the claim that Arie's SO2 notation meant sulfur-dioxide release, biological sulfur signaling, or a coded energetic/quantum mechanism.",
    "The patent does not establish membrane separation, ion exchange, biocompatibility, blood compatibility, cell response, or medical-device suitability.",
    "The corpus does not establish universal constant-time computation, zero memory, zero energy, exact arbitrary long-horizon prediction, solved general continuous collision detection, or elimination of all broad-phase indexing.",
    "No historical source demonstrates that Arie intended Tom's later topological vocabulary. The overlap is a present-day comparative interpretation, not a claim of concealed lineage.",
])

# ---------------------------------------------------------------------------
# 2 Arie
# ---------------------------------------------------------------------------
b.heading("2. Arie Klootwijk's technical core",1,page_break=True)
b.para("Arie's patent is best understood as a topology-control process for a reactive engineering thermoplastic. Its preferred feed pair is 4,4'-sulfonyldiphenol - now commonly called bisphenol S (BPS) - and its diglycidyl ether (DGE-BPS). A phenoxide attacks an epoxide ring, creating an aryl-ether bond and a secondary hydroxyl. Repetition produces an almost exclusively linear poly(hydroxy ether) containing aromatic sulfone bridges [S1, pp. 3-5; S2].",style="Lead")

# two-column image/text using table
imgtable=b.doc.add_table(rows=1,cols=2)
imgtable.alignment=WD_TABLE_ALIGNMENT.CENTER
imgtable.autofit=False
c1,c2=imgtable.rows[0].cells
set_cell_width(c1,6.0);set_cell_width(c2,10.0)
set_cell_margins(c1,top=50,start=50,bottom=50,end=100);set_cell_margins(c2,top=50,start=120,bottom=50,end=50)
p=c1.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(FIG/"source_patent_front_page.png"),width=Cm(5.3))
p2=c2.paragraphs[0]
r=p2.add_run("HISTORICAL SOURCE")
r.bold=True;r.font.color.rgb=rgb(CORAL);r.font.size=Pt(9)
p3=c2.add_paragraph();add_run_with_citations(p3,"The Swedish front page identifies Shell Internationale Research Maatschappij N.V. as applicant and R. W. Kreps, A. Klootwijk, and J. M. Goppel as inventors. The title states a method for thermoplastics from a dihydroxy and diepoxy compound in which the diepoxy compound or both compounds contain an SO2 group. Priority is claimed from the Netherlands on 15 June 1959 [S1, p. 1].",size=9.2)
p4=c2.add_paragraph();p4.style=b.doc.styles["Small"];p4.add_run("The patent is a process document, not merely a composition sketch. Its novelty is the integrated control of stoichiometry, solvent reactivity/polarity, solubility, purity, water, catalyst, and endpoint.")
set_table_no_outer_borders(imgtable)
b.md.append("![Swedish patent front page](../03_Figures/source_patent_front_page.png)\n\n")

b.figure(FIG/"figure_00_patent_reaction_concept.png","Figure 2. Preferred chemistry in the patent: BPS plus DGE-BPS gives a sulfone-containing linear poly(hydroxy ether). Each epoxide-opening event forms an ether bond and a pendant secondary hydroxyl.",width_cm=16.5)

b.heading("2.1 The SO2 group: what it is and what it does",2)
b.para("The printed SO2 unit is the neutral covalent sulfone bridge Ar-S(=O)2-Ar. It is not free sulfur dioxide, and it is not a sulfonate ion. In the preferred BPS/DGE-BPS system, the sulfone joins two aromatic rings. Its high polarity and rigid aromatic environment affect chain stiffness, dipole interactions, solubility, packing, glass/thermal behavior, water affinity, and adhesion. The epoxide-derived hydroxypropyl segments add regularly spaced secondary hydroxyls that can hydrogen-bond or be derivatized.")
b.cards([
    ("NOT SULFUR DIOXIDE", "No SO2 gas is generated by the repeat-unit notation. The sulfur is part of a stable covalent diaryl-sulfone bridge.", CORAL),
    ("NOT AN ION-EXCHANGE GROUP", "A neutral sulfone does not provide fixed ionic charge. Ion-exchange behavior would require sulfonate, ammonium, phosphonate, or other ionizable groups.", GOLD),
    ("A POLAR STRUCTURAL BRIDGE", "The sulfone raises polarity and rigidity while remaining part of the backbone architecture.", TEAL),
    ("AN INTERFACE-READY POLYMER", "The newly created secondary hydroxyls are the most useful post-functionalization handles for coatings, blends, crosslinking, or immobilization.", GREEN),
],cols=2)

b.heading("2.2 Equimolarity as an invariant of chain growth",2)
b.para("The phrase 'substantially equimolar' is the mathematical center of the patent. In a bifunctional AA + BB step-growth process, high molecular weight appears only when complementary functional groups are almost exactly balanced and conversion approaches unity. A small imbalance creates a hard ceiling even if every limiting group reacts. The ideal Carothers-Flory relation is:")
b.equation("X_n = (1 + r) / (1 + r - 2 r p)      and at p = 1: X_n,max = (1 + r) / (1 - r)","r is the ratio of limiting to excess functional groups; p is conversion of the limiting group.")
b.para("At a one-percent imbalance, the ideal ceiling is roughly two hundred repeat-unit equivalents; at two percent it is roughly one hundred; at three percent it is about sixty-six. Exact balance is still insufficient unless conversion is extremely high: at r = 1, 99% conversion gives Xn about 100, 99.5% gives about 200, and 99.9% gives about 1,000. This is why solvent and phase control are inseparable from the ratio claim.")
b.figure(FIG/"figure_03_stoichiometry_as_invariant.png","Figure 3. Ideal step-growth theory shows that equimolarity and conversion jointly control attainable chain length. The curve is a quantitative conservation constraint, not a topological visual pun.",width_cm=15.5)

b.heading("2.3 The subtle catalyst bookkeeping",2)
b.para("The patent explicitly counts dihydric phenol present as an alkali-metal salt in the dihydroxy:diepoxy ratio. Example 1 charges 1.00 mol DGE-BPS, 0.97 mol neutral BPS, and 0.05 mol monosodium BPS. On the patent's own accounting, the effective BPS-derived amount is 1.02 mol, exactly the phenol-rich edge of the preferred window. The salt is catalytic because phenoxide character can be transferred, but its BPS skeleton can also enter the material balance. This is unusually careful process logic: the catalyst is not treated as informationally invisible.")

b.heading("2.4 Solvent as an accessibility field",2)
b.para("The invention arose because prior solvent choices allowed the product to crystallize or acted as chain terminators. Arie's solution is not simply 'use a better solvent'. The solvent must simultaneously dissolve the dihydroxy compound, the diepoxide, and the growing polymer; have sufficient polarity; scarcely react with phenol or epoxide; avoid catalyzing side reactions; and remain liquid at the reaction temperature. Ketones, nitriles, nitro compounds, sulfoxides, and sulfones are listed, with dielectric constant used as a practical proxy [S1, pp. 3-4].")
b.callout("Topological reading", "Solvent quality determines whether complementary chain ends remain in the active connected component of reaction space. Premature crystallization is a physical removal of growing chains from the accessible relation set. This is a valid structural analogy to support/reachability - but the governing law is solution thermodynamics and transport, not a Boolean cone query.", color=GOLD, fill=LIGHT_GOLD)

b.heading("2.5 Molecular topology actually controlled",2)
b.table(
    ["Control variable","Physical effect","Topology consequence","Failure if uncontrolled"],
    [
        ["Difunctional monomers","Two phenolic and two epoxide functions","Favors linear AA+BB chains","Multifunctional impurities create branching/networking"],
        ["Near-equimolarity","Balances complementary chain ends","Raises attainable chain length; programs end groups","Small imbalance caps molecular weight"],
        ["High conversion","Consumes remaining functional ends","Connects oligomers into long chains","Low conversion leaves short disconnected chains"],
        ["Nonreactive polar solvent","Maintains mobility and suppresses chain transfer","Keeps reactive ends mutually reachable","Alcohol/water can cap chains; poor solvent precipitates product"],
        ["Monomer purity","Limits monofunctional and geometric defects","Reduces branch points and packing defects","o,p'-isomer, chloride, glycols, or mono-functions change architecture"],
        ["Temperature/oxygen control","Controls rate and degradation","Preserves intended edge-creation rule","Side reactions, oxidation, scission, or uncontrolled branching"],
        ["Chain stopper/precipitation","Closes the process at chosen state","Sets end groups and freezes distribution","Continued reaction changes viscosity during processing"],
    ],
    widths=[3.0,4.0,4.2,4.8],font_size=7.8
)

b.heading("2.6 What the experimental examples establish",2)
b.para("The examples show that reaction time and phase control move the product across a sharp property threshold. Example 1 reaches intrinsic viscosity 0.43 and reports strong pressed specimens, chemical resistance, and a heat-distortion temperature around 154 C under the historical test. Example 2, stopped after 20 hours at lower intrinsic viscosity 0.35, has much lower impact performance. Example 3 re-dissolves crystallizing product at higher temperature and reaches 0.45-0.50, restoring strong mechanical behavior. Example 4 rapidly reaches high viscosity and deliberately adds BPS as a chain stopper. Example 5 shows sulfolane as a suitable solvent. Example 6 substitutes a bisphenol-A diepoxide while retaining BPS as the sulfone-containing dihydroxy component [S1, pp. 6-8].")
b.para("These data are historically meaningful but incomplete by modern standards. Intrinsic viscosity is not an absolute molecular-weight distribution; the examples do not report SEC/MALS, NMR end groups, branching, residual monomer, morphology, fracture statistics, long-term aging, or extractables. The correct conclusion is not that the material is fully characterized, but that the process produced a credible high-molecular-weight engineering thermoplastic and demonstrated why maintaining accessibility to high conversion mattered.")

b.heading("2.7 The patent as an information process",2)
b.para("Read abstractly, the process carries information in matter. The feed ratio encodes attainable chain length and end-group bias. Monomer purity constrains the allowed rewrite grammar. Solvent determines reachability. Temperature schedules control event rate. Intrinsic viscosity is a coarse state estimator. Precipitation and chain stopping commit the batch. Each pendant hydroxyl is a repeated interface address for later chemistry. This is not proof that the patent hid a computer, but it shows why the comparison is technically fertile: polymer synthesis is a physical graph transformation with a distributed state and an irreversible log embodied in the product.")

# ---------------------------------------------------------------------------
# 3 Tom
# ---------------------------------------------------------------------------
b.heading("3. Tom Klootwijk's mature technical core",1,page_break=True)
b.para("The supplied corpus begins in graphics-adjacent language - cones, SDFs, polar coordinates, one-bit fields, double vacuum, Klein bottles, hourglasses, eigenvectors, and zero-crossings - but its strongest internal development is a correction away from treating those motifs as rendering primitives or literal physics. The mature interpretation is an ontology of directly queryable relations, events, phases, identities, and local supports. Projection into images is optional [S3, pp. 1, 5-10].",style="Lead")

b.heading("3.1 Formal core",2)
b.para("A compact version of the proposed state description is a product space containing physical position, continuous time, phase, a sheet/parity label, a lineage address, and a branch. Relations define event surfaces; support predicates define relevance; compatibility predicates decide whether co-located sectors may interact; and transition operators update state while preserving declared invariants.")
b.equation("q_e(t) = (x_e(t), t, phi_e(t), sigma_e(t), a_e, b_e)","Entity state: position, time, phase, sheet/orientation, address, branch.")
b.equation("t* = inf { t >= t0 : R_j(q_e(t)) = 0,  C_alpha(q,t) <= 0,  chi(e,j,t) = 1 }", "Earliest admissible event under relation, local support, and compatibility.")
b.equation("q_e(t*+) = T_j(q_e(t*-), context)","Transition updates the state; lineage/invariants and exogenous novelty are logged.")

b.heading("3.2 Local spherical support, not a spherical universe",2)
b.para("The mature corpus repeatedly narrows 'spherical' to a local chart around an agent, sensor, coupler, shell, or region of influence. Radial reach and angular orientation are native for field of view, support, uncertainty, and coupling. They need not replace the global world representation. This is a major strength because local radial-angular coordinates are already natural in sensing, scattering, acoustics, optics, robotics, and biointerfaces [S3, pp. 5-7; S6, pp. 3-4].")

b.heading("3.3 Double vacuum as absent coupling",2)
b.para("In the disciplined reading, two states may share the same coordinate yet remain mutually invisible because phase, sheet, orientation, mode, address, provenance, or policy is incompatible. The 'second vacuum' is therefore not a deeper physical emptiness; it is the absence of an allowed coupling. This translates cleanly into orthogonal photonic modes, separate microfluidic channels, receptor mismatch, security domains, typed states, or disconnected phases. It becomes useful precisely when the compatibility schema is explicit.")

b.heading("3.4 SDF zero as event semantics",2)
b.para("Earlier material was tempted to make SDF = 0 a way to move 'past' rasterization by continuing a sampled field. The mature synthesis instead treats zero as an explicit transition relation: the point at which event semantics, routing, confidence, and lineage are invoked. This aligns the corpus with implicit surfaces and hybrid systems rather than with a magical escape from numerical computation [S3, pp. 7-10].")

b.heading("3.5 One-bit parity in a narrow role",2)
b.para("One-bit is most defensible as a parity, route, validity, freshness, availability, or compatible/incompatible flag. Spherical Throughput is explicit that optical amplitudes, thresholds, uncertainty, and lineage remain separate state. This is essential. A binary event can summarize a decision, but it cannot represent the physical process that produced the decision [S6, p. 7].")

b.heading("3.6 The most mature engineering translation: B.C.E.",2)
b.para("The Bounded Compatibility Event is the corpus's clearest bridge from metaphor to engineering. An output is counted only when it lies in declared support, matches the selected channel, crosses a measured guard, and meets a confidence or certification threshold. The proposed throughput is therefore verified events per second at a declared error budget, not raw photon flux, pixel rate, or rhetoric [S6, pp. 3, 6-7].")
b.equation("N_verified = sum_q 1[S_q] 1[chi_q] 1[g_q crossed] 1[c_q >= c_min]", "A B.C.E. requires support, compatibility, a measured crossing, and confidence.")

b.heading("3.7 Feasible core versus exploratory overclaim",2)
b.table(
    ["Corpus claim or motif","Retain / rewrite / reject","Disciplined interpretation"],
    [
        ["Direct state-at-time query","Retain with bounds","Can be horizon-independent for a fixed closed expression; still depends on expression size, branch history, and numerical conditioning."],
        ["Next event without frame replay","Retain with bounds","Solve or conservatively bound roots for restricted relation families; event density and degeneracies remain costs."],
        ["Local spherical support","Retain","Use radial-angular domains for relevance, sensing, coupling, and uncertainty."],
        ["Double vacuum","Translate","Use typed, phase-separated, orthogonal, uncoupled, or policy-isolated sectors."],
        ["One-bit world","Rewrite","Use a bit as route/validity/parity; maintain continuous and structured state separately."],
        ["Klein-bottle hardware","Demote to model","Require an explicit gluing/routing map; literal non-orientable fabrication is a frontier, not a premise."],
        ["Universal O(1), zero memory","Reject","Exogenous novelty, candidate relations, branches, event logs, and numerical work accumulate."],
        ["Solved general CCD / no sieves","Reject as universal","Analytic time-of-impact can help restricted shapes; broad-phase/support/indexing is still needed when candidate sets grow."],
        ["World equals general AI","Reject as established","A shared state substrate may aid embodied agents, but does not collapse arbitrary learning into a small transform matrix."],
    ],
    widths=[4.1,3.1,8.8],font_size=8.0
)

b.figure(FIG/"figure_04_evidence_readiness.png","Figure 4. Arie and Tom have complementary strengths. The scores are qualitative and compare evidence/readiness within this review, not intrinsic intellectual merit.",width_cm=15.4)

b.heading("3.8 Where Tom's work is genuinely ahead",2)
b.bullets([
    "It makes compatibility a first-class condition instead of assuming that spatial proximity implies interaction.",
    "It separates authoritative state, local support/sensing, and downstream projection.",
    "It treats identity as lineage plus invariants rather than only instantaneous coordinates.",
    "It can integrate heterogeneous sensors and models through local-to-global consistency methods such as sheaves [R04].",
    "It naturally fits hybrid systems: continuous dynamics punctuated by guarded transitions [R19].",
    "It provides a native place for uncertainty, calibration, false/missed-event rates, and event replay.",
    "It can become an event-sourced digital material or biological interface, a capability unavailable in a 1959 process patent.",
])

# ---------------------------------------------------------------------------
# 4 Four topology layers
# ---------------------------------------------------------------------------
b.heading("4. Four layers of topology",1,page_break=True)
b.para("The central analytical result is that the two bodies of work are strongest at opposite ends of a four-layer stack. Arie directly controls molecular graph topology. Tom directly describes operational and information topology. Between them lie morphology/phase topology and interface/transport topology - the layers where membranes, biointerfaces, photonics, and digital twins can connect the work.",style="Lead")
b.figure(FIG/"figure_01_four_topology_layers.png","Figure 5. Four distinct topology layers. Productive cross-domain work connects the layers with explicit models and measurements rather than treating them as synonyms.",width_cm=15.8)

topology_rows=load_csv(DATA/"topology_layers.csv")
b.table(
    ["Layer","Arie's direct contribution","Tom's direct contribution","Methods that connect them"],
    [[r["Layer"],r["Arie"],r["Tom"],r["Useful_methods"]] for r in topology_rows],
    widths=[3.2,4.5,4.5,3.8],font_size=7.6
)

b.heading("4.1 Molecular graph topology",2)
b.para("At the molecular layer, topology means the graph of atoms and covalent bonds, together with functionality, branch points, cycles, connected components, and end groups. Arie's process changes this graph directly. Difunctionality favors degree-two chain interiors; stoichiometric imbalance changes the number and identity of terminal vertices; monofunctional impurities terminate components; multifunctional impurities introduce branch points; intramolecular reactions can create cycles; and crosslinking can percolate into a network. This is graph topology in a literal chemical sense.")
b.para("Tom's graph-rewriting language can model these operations, and chemical graph grammar is an established field [R18]. But the computational rule must be attached to reaction rates, concentrations, sterics, solvent, temperature, and mass balance. Without those, it is a symbolic catalogue of possible rewrites rather than a predictive polymerization model.")

b.heading("4.2 Morphology and phase topology",2)
b.para("After covalent synthesis, the material has another topology: crystalline and amorphous domains, chain entanglement, free volume, pores, and interfaces. Arie's patent confronts this layer during polymerization because crystallization can remove product from solution and stop chain growth. In a membrane, the same layer determines connected pores, tortuosity, dead ends, percolation, and the evolving topology of fouling or biofilm.")
b.para("Tom's phase sheets and double-vacuum sectors become physically meaningful here only when mapped to measured phases or transport-isolated domains. Persistent homology can characterize components, loops, and voids across a filtration scale and has been applied to pore geometry [R05]. Reeb graphs, merge trees, and contour trees can summarize time-varying scalar fields in scientific visualization [R06].")

b.heading("4.3 Interface and transport topology",2)
b.para("This is the most promising joint layer. The polymer's pendant hydroxyls and polar backbone determine which molecules, supports, coatings, fillers, proteins, cells, or optical modes can couple. Tom's local support and compatibility predicates determine which potential interactions are admitted. A membrane pore may be spatially open yet electrostatically incompatible; a receptor and analyte may be co-located yet chemically mismatched; two waveguide modes may overlap geometrically yet be orthogonal; a cell may contact a surface yet fail to form stable adhesion.")
b.para("The correct model is multiplicative rather than metaphorical: interaction requires spatial support, transport access, chemical compatibility, and sufficient kinetics. One possible factorization is chi_total = chi_geometry * chi_transport * chi_chemistry * chi_sensor * chi_policy, with each term carrying a probability or uncertainty rather than automatically being a perfect bit.")

b.heading("4.4 Operational and information topology",2)
b.para("Tom is strongest here. A membrane module, biointerface chip, or material coupon has an operational topology consisting of sensor neighborhoods, calibration relationships, admissible states, event guards, transitions, and lineage. Sheaf methods can test whether local observations agree globally; hybrid automata can represent continuous flux or cell motion with discrete cleaning or adhesion events; event sourcing can reconstruct the state from an append-only history; and digital-twin methods can update hidden states and uncertainty from data [R04, R07, R19-R20].")

b.callout("A five-layer implementation stack", "For a serious prototype, implement (1) molecular/reaction graph, (2) morphology and phase state, (3) transport/interface model, (4) measurement and uncertainty model, and (5) operational event/lineage layer. Tom's calculus belongs primarily in layer 5 and can coordinate the others; it should not erase them.", color=TEAL, fill=LIGHT)

# ---------------------------------------------------------------------------
# 5 Crosswalk
# ---------------------------------------------------------------------------
b.heading("5. Detailed fundamental crosswalk",1,page_break=True)
b.para("The comparison becomes most informative when each corpus term is matched not to a superficial visual resemblance, but to a role in a constrained transformation system. Figure 6 presents the shared grammar; the sections below explain where each mapping holds and where it breaks.",style="Lead")
b.figure(FIG/"figure_02_event_grammar_crosswalk.png","Figure 6. Shared event grammar. The top and bottom lanes are not identical mechanisms; they instantiate the same abstract sequence at different layers.",width_cm=16.5)

b.heading("5.1 Support: solvent mobility versus local relevance",2)
b.para("Arie's support is physical. A chain end participates only if the molecule is dissolved, mobile, chemically active, and able to encounter a complementary function. Tom's support is an analytic or semantic domain: a cone, shell, hourglass, field of view, causally reachable set, or local query neighborhood. The genuine overlap is that presence alone is insufficient; accessibility must be declared. The difference is that physical support is governed by diffusion, concentration, viscosity, phase equilibrium, and geometry.")

b.heading("5.2 Compatibility: stoichiometry versus typed coupling",2)
b.para("At each chemical event, a nucleophilic phenoxide and an epoxide are complementary types. Across the batch, the total numbers of those functions must remain nearly balanced if long chains are to form. Tom's chi predicate generalizes type matching to phase, sheet, address, mode, policy, provenance, or time window. The shared principle is selective coupling. The critical distinction is scale: equimolarity is a global conservation constraint over a population, whereas chi is normally a local admission condition.")

b.heading("5.3 Event: ring opening versus guard crossing",2)
b.para("A ring-opening event creates a new bond, destroys an epoxide, creates an alcohol, changes local charge/protonation transiently, and alters the chain-length distribution. A computational guard crossing changes discrete state according to a transition rule. Both are local rewrites. Chemical events, however, are probabilistic and embedded in competing pathways. A robust digital twin should therefore treat reaction or transport guards as stochastic/interval events rather than exact algebraic instants unless measurement and solver certify the crossing.")

b.heading("5.4 Transition: covalent connectivity versus routing/state update",2)
b.para("Arie's transition physically merges molecular components. Tom's transition routes a state to another sheet, chamber, mode, branch, or semantic condition. In a membrane, both types can coexist: a surface reaction changes covalent chemistry while an operational transition changes the module from production to cleaning. A multilayer event record should distinguish the material transition from the controller transition.")

b.heading("5.5 Lineage and closure",2)
b.para("A polymer batch carries process history in its molecular-weight distribution, end groups, residuals, morphology, and thermal/mechanical state. Tom makes this history explicit through lineage addresses, invariants, and an event log. The natural present-day synthesis is a digital material passport that links raw-material lots, stoichiometric calculations, catalyst/solvent, temperature history, isolation, functionalization, sterilization, aging, calibration, and device events. This is more faithful than assigning a coordinate-derived UUID to an individual polymer chain.")

b.heading("5.6 The striking analogy of premature closure",2)
b.callout("Generational through-line", "Arie discovered that premature crystallization or chain termination prevented the system from reaching useful molecular weight. Tom's mature corpus argues that premature projection into frames, pixels, voxels, or conventional object inventories can prevent the system from preserving its query-first relational structure. In both cases, the medium must keep relevant relations available until the high-value closure condition is reached. This is an analogy of process architecture, not evidence of hidden historical transmission.", color=GOLD, fill=LIGHT_GOLD)

b.heading("5.7 Compact comparison matrix",2)
comp=load_csv(DATA/"comparison_matrix.csv")
# split into two tables for readability
for chunk_no in range(0,len(comp),7):
    chunk=comp[chunk_no:chunk_no+7]
    b.table(
        ["Dimension","Arie","Tom","Comparative verdict"],
        [[r["Dimension"],r["Arie_1959_technique"],r["Tom_current_technique"],r["Genuine_overlap"]+" Difference: "+r["Critical_difference"]] for r in chunk],
        widths=[2.7,4.3,4.3,4.7],font_size=7.25
    )

b.heading("5.8 Where analogy must stop",2)
b.bullets([
    "A covalent bond is not a database edge unless the model explicitly maps and measures it.",
    "A physical phase is not an abstract phase angle; use separate variables such as phase_phys and phase_info.",
    "A solvent dielectric constant is not a universal relevance score.",
    "A viscosity threshold is an empirical state proxy, not an exact molecular topology measurement.",
    "A one-bit flag cannot encode continuous composition, chain distributions, morphology, or biological state.",
    "A non-orientable surface analogy does not establish non-orientable material geometry or transport.",
    "Closed-form notation does not guarantee closed-form solvability, numerical stability, or constant cost.",
])

# ---------------------------------------------------------------------------
# 6 CS/CG
# ---------------------------------------------------------------------------
b.heading("6. Topological information science and computer graphics",1)
b.para("Tom's work is most defensible when translated into a hybrid symbolic-numeric architecture for state and event queries. Its novelty is less 'sphere instead of grid' than 'query and transition semantics instead of treating projection as the world'. That position is compatible with existing graphics and information-science methods while still leaving a distinct research program.",style="Lead")

b.heading("6.1 Translation into established formalisms",2)
b.table(
    ["Corpus term","Closest established formalism","Useful implementation","Caution"],
    [
        ["Local sphere / cone","Domain of dependence, kernel support, sensor frustum, influence volume","Analytic support predicate; radial-angular indexing; uncertainty cone","Do not remesh the entire world into polar coordinates"],
        ["SDF = 0 / B = 0","Implicit surface, level set, hybrid guard","Root isolation, conservative advancement, interval arithmetic, event detection","An arbitrary SDF intersection may still require numerical iteration"],
        ["Double vacuum","Typed channel, mode orthogonality, covering-space sheet, hidden state","Compatibility label and transfer/coupling matrix","Same coordinate does not prove meaningful separate physical space"],
        ["One-bit","Parity, route, validity, orientation, freshness","Bit mask accompanying richer state","Do not collapse amplitude, confidence, or provenance into the bit"],
        ["Hourglass / quad chambers","Double cone, causal/support partition, branching/routing state","Finite routing automaton or stratified state space","Visual symmetry does not guarantee algebraic closure"],
        ["Ontological UUID","Persistent entity identifier, generative address, event-sourced aggregate","Stable ID plus split/merge lineage graph","Coordinate should not be identity when entities move or merge"],
        ["Equation world","Hybrid system, procedural implicit scene, symbolic dynamics","State evaluator, event solver, transition router, log","General scenes still need indexing, approximation, and data for novelty"],
        ["No sieves","Support/compatibility pruning","Analytic culling and typed candidate reduction","Cannot remove candidate-selection cost for arbitrary large relation sets"],
    ],
    widths=[2.7,4.0,4.7,4.6],font_size=7.5
)

b.heading("6.2 Representation versus projection",2)
b.para("Rasterization asks which stored surfaces project into pixels now. Ray marching asks what repeated field samples reveal along a projected direction. Ray tracing asks which transport paths connect sources, surfaces, and sensors. Tom's mature substrate asks which relations generate possible state, what state holds at time t, what admissible event occurs next, and which sectors may couple. Images can still be emitted, but they are downstream views [S3, pp. 9-10].")
b.para("This is a meaningful inversion for simulation, digital twins, and scientific visualization. It resembles procedural implicit geometry, event-driven simulation, temporal databases, and hybrid automata. The hard research question is not whether it sounds different from rendering; it is whether support pruning plus event solving costs less than the state, frames, or candidate structure it avoids materializing.")

b.heading("6.3 Restricted Equation World Zero",2)
b.para("The correct first prototype is intentionally small: a two-dimensional homogeneous/projective state, continuous time, two sheets and a routing bit, a bounded relation family, simple trajectory classes, finite grammar depth, lineage-based identity, and symbolic state/event outputs. The prototype should answer six queries: state at time, next event, events in support, phase coupling of co-located sheets, transition routing, and state reconstruction from seed plus event log [S3, pp. 7, 11-12].")
b.table(
    ["Experiment","Success evidence","Failure mode to expose"],
    [
        ["Horizon skipping","Cost follows expression/branch complexity rather than skipped frame count","Expression expansion or numerical conditioning grows with horizon"],
        ["Co-located sheets","Incompatible states remain uncoupled until chi becomes true","Sheet labels merely duplicate coordinates without meaningful semantics"],
        ["Event ordering","Conservative/exact solver returns correct first event and stable transition","Tangencies, multiple roots, or degeneracy produce missed/reordered events"],
        ["Grammar depth","Relations remain normalized or bounded under composition","Branch/expression explosion erases direct-query advantage"],
        ["Identity split/merge","Lineage remains reconstructable and collision-safe","Coordinate-derived identity breaks under merge, split, or reconciliation"],
        ["Matched baseline","Declared workload beats frame stepping/BVH or explains where it does not","No performance or correctness advantage"],
    ],
    widths=[3.4,6.2,6.4],font_size=8.0
)

b.heading("6.4 Sheaves: the natural mathematics of local compatibility",2)
b.para("Sheaf theory is a particularly strong external match to the corpus. A sheaf assigns data to local regions and defines restriction maps between overlapping regions; global consistency exists when local assignments agree on overlaps. This directly formalizes a world where sensors, agents, models, phases, or interface patches hold local state and must be reconciled without forcing one monolithic raster [R04].")
b.para("For membranes or biointerfaces, a sheaf can organize local pressure, optical, chemical, electrical, and biological measurements. Inconsistency becomes measurable rather than rhetorical: two local states may be individually plausible yet fail to glue globally. That is a mathematically disciplined version of the corpus's compatibility sectors and double-vacuum intuition.")

b.heading("6.5 Persistent topology and scientific visualization",2)
b.para("Persistent homology and related descriptors provide a rigorous route from the corpus's topology language to measurable geometry. A filtration converts image intensity, distance, concentration, or threshold into a family of spaces. Components, loops, and voids that persist across scale are more robust than a single threshold. In porous media, persistence descriptors have been used to characterize pore heterogeneity [R05]. In computer graphics and visualization, persistence diagrams, merge trees, contour trees, Reeb graphs, and Morse-Smale complexes support scalar-field comparison across single fields, time series, and ensembles [R06].")
b.para("This is relevant to Arie's material because morphology may evolve during polymerization, precipitation, membrane formation, fouling, or biological colonization. It is relevant to Tom because topology becomes a measured descriptor and query target rather than a shape metaphor.")

b.heading("6.6 Event-based sensing",2)
b.para("Event cameras report asynchronous intensity changes instead of global frames. They therefore instantiate the corpus's 'event, not frame' preference at the sensor level, with well-understood limitations. Event-based sensing has been demonstrated for high-speed particle tracking in microfluidic devices [R08] and for event-based imaging flow cytometry combined with photonic neuromorphic processing [R09]. These precedents do not validate the whole equation-world architecture, but they make sparse bio/microfluidic event acquisition a credible near-term application.")

b.heading("6.7 Graph grammars and chemical rewriting",2)
b.para("Graph grammar provides the cleanest formal bridge back to Arie's work. Molecules are graphs; elementary reactions are local rewrites; reaction networks can be generated by composing rules [R18]. For polymerization, a grammar can express epoxide ring opening, chain merging, branch creation, termination, or side reactions. Tom's finite grammar can therefore be grounded in chemical graph rewriting - but predictive use still requires rates, populations, conservation laws, and transport.")

b.heading("6.8 Where Tom's work is best positioned in CS/CG now",2)
b.cards([
    ("QUERY-FIRST SCIENTIFIC DIGITAL TWINS", "Direct state/event queries over reduced physical models, with sparse updates, uncertainty, and lineage.", TEAL),
    ("TOPOLOGY-AWARE INTERFACE ANALYTICS", "Persistent descriptors and Reeb/merge structures for pores, fouling, deformation, or cell morphology.", PURPLE),
    ("EVENT-BASED MICROFLUIDIC VISION", "Asynchronous particle/cell events, support gating, and reference-frame validation.", GOLD),
    ("RESTRICTED IMPLICIT EVENT SOLVERS", "Analytic or certified root solving for controlled relation families, benchmarked against conventional CCD and spatial indices.", CORAL),
    ("MATERIAL-AWARE PROCEDURAL GRAMMARS", "Generative geometry and chemistry whose rules carry material, transport, and lineage semantics rather than only polygons.", GREEN),
    ("LOCAL-TO-GLOBAL SENSOR FUSION", "Sheaf-style compatibility checks across heterogeneous local charts, sensors, and agents.", NAVY),
],cols=2)

# ---------------------------------------------------------------------------
# 7 Membranes and biointerfaces
# ---------------------------------------------------------------------------
b.heading("7. Membranes and biointerfaces",1,page_break=True)
b.para("The strongest cross-generational application is a material interface whose chemistry is physically real and whose operation is natively queryable. Arie's platform supplies a hydroxyl-rich, sulfone-containing thermoplastic architecture. Tom supplies local support, compatibility, bounded events, uncertainty, topology descriptors, and lineage. The combined system should be described as a chemically addressable interface with an event-sourced twin - not as a mystical topological membrane.",style="Lead")

b.heading("7.1 Material roles for a PHES/phenoxy-like polymer",2)
b.table(
    ["Role","Why the chemistry fits","What must be measured"],
    [
        ["Thin reactive coating","OH groups support grafting/crosslinking; sulfone and aromatic backbone give polarity and heat resistance","Coating thickness, coverage, adhesion, swelling, permeability penalty, defects, residuals"],
        ["Tie layer in thin-film composite","Hydrogen bonding and derivatization can couple support and selective layer","Interfacial fracture, delamination, solvent/cleaning stability, transport resistance"],
        ["Blend modifier","High-MW thermoplastic can alter toughness, compatibility, and phase morphology","Blend miscibility, phase topology, pore formation, Tg, modulus, water uptake"],
        ["Functionalization scaffold","Repeated secondary OH groups provide addresses for PEG, zwitterions, polysaccharides, peptides, ligands, or crosslinkers","Degree/distribution of functionalization, ligand activity, leachables, aging"],
        ["Patterned microfluidic/biosensor surface","Thermoplastic processing and surface chemistry support microdevices and immobilization","Optical background, nonspecific adsorption, channel bonding, sterilization, sensor drift"],
    ],
    widths=[3.5,6.2,6.3],font_size=8.0
)

b.heading("7.2 What the material is not",2)
b.bullets([
    "It is not commercial polysulfone or polyethersulfone; the regular hydroxypropyl segments and pendant OH groups make it a different, more chemically addressable poly(hydroxy ether).",
    "It is not an ion-exchange membrane solely because it contains sulfone. The sulfone is neutral.",
    "It is not automatically antifouling. Hydroxyls and polarity can increase water affinity but may also support protein interactions; surface composition and hydration must be measured.",
    "It is not automatically biocompatible. Residual monomers, oligomers, catalyst, solvent, functionalization chemistry, sterilization, degradation, and intended contact category control the biological-risk assessment.",
])

b.heading("7.3 Membrane architecture options",2)
b.cards([
    ("A. SUPPORT + REACTIVE COATING", "Start with a known porous support. Apply a thin hydroxyl-rich layer, then derivatize or crosslink. Lowest-risk path because mechanics and pores are supplied by an established membrane.", TEAL),
    ("B. BLEND / PHASE-INVERSION ADDITIVE", "Blend the polymer or analogue into a membrane-forming matrix. Higher leverage over morphology but greater risk of phase separation, leaching, and pore/flux changes.", GOLD),
    ("C. THIN-FILM COMPOSITE INTERLAYER", "Use the polymer as an adhesive, toughening, or reactive interlayer under a selective skin. Strong fit to its tie-layer character.", CORAL),
    ("D. AFFINITY OR BIOACTIVE SURFACE", "Functionalize OH groups with a ligand, zwitterion, anticoagulant, peptide, or capture reagent. Highest biological specificity and highest safety/validation burden.", PURPLE),
],cols=2)

b.heading("7.4 Membrane World Zero",2)
b.figure(FIG/"figure_05_membrane_world_zero.png","Figure 7. Membrane World Zero couples material lineage and continuous physical state to B.C.E. guards and an append-only event store. The material, physical model, sensors, estimator, and control remain distinct.",width_cm=16.4)
b.para("The minimum demonstrator should use a commercial baseline membrane before introducing the inherited chemistry. Pressure, cross-flow, permeate flux, conductivity, temperature, and optional optical fouling are logged conventionally and through an event layer. A hidden-state estimator tracks membrane resistance and fouling resistance with uncertainty. Guards emit events for flux decline, pressure rise, breakthrough, cleaning recovery, topology change, and calibration drift. The event representation is successful only if it improves decision latency, bandwidth, energy, or interpretability at equal false/missed-event limits.")
b.table(
    ["State / guard","Physical meaning","Event-layer role"],
    [
        ["Flux J and TMP","Transport performance and driving force","Continuous state; slopes/levels define candidate fouling or blockage guards"],
        ["Fouling resistance Rf","Hidden or inferred accumulated resistance","Digital-twin state with uncertainty; updates from measurements"],
        ["Conductivity/tracer","Selectivity or breakthrough","Compatibility and threshold event"],
        ["Optical topology","Cake, biofilm, bubbles, or pore morphology","Persistent descriptors can detect structural regime changes"],
        ["Calibration residual","Sensor/model trustworthiness","Invalidates or downgrades events; prevents a one-bit false certainty"],
        ["Material lineage","Batch, coating, graft, sterilization, cycles","Allows replay and links performance to chemistry rather than only geometry"],
    ],
    widths=[3.4,6.1,6.5],font_size=8.1
)

b.heading("7.5 Biointerface World Zero",2)
b.figure(FIG/"figure_06_biointerface_world_zero.png","Figure 8. Biointerface World Zero: surface chemistry is linked to sparse interaction events, topology descriptors, confidence, and replayable lineage.",width_cm=16.0)
b.para("A first biointerface study should compare an inert control, an hydroxyl-rich sulfone polyether or safer analogue, an antifouling derivatization, and an affinity or cell-adhesive derivatization. Event imaging records fast attachment, motion, aggregation, and detachment, while periodic conventional images remain the reference. Persistent topology summarizes connected clusters, loops, and voids; local optical, impedance, or chemical measurements are fused with consistency checks.")
b.table(
    ["Interface function","Candidate chemistry","Readout and event"],
    [
        ["Antifouling","PEG, zwitterionic, hydrated polysaccharide, or dense neutral brush","Protein adsorption; first stable adhesion; cluster count; detachment; long-term drift"],
        ["Affinity capture","Covalently immobilized ligand, antibody fragment, aptamer, peptide, or chelator","Binding onset, occupancy, breakthrough, regeneration, false/missed capture"],
        ["Cell-adhesive","Peptide or extracellular-matrix motif at controlled density","Attachment, spreading, migration, proliferation, detachment force"],
        ["Blood-contact research","Hydrophilic/zwitterionic or anticoagulant surface","Platelet adhesion/activation, coagulation, complement, hemolysis; application-specific standards"],
        ["Optical/electrical biosensor","Immobilized receptor plus waveguide, fluorescence, impedance, or field-effect readout","Compatibility-gated threshold with confidence and calibration lineage"],
    ],
    widths=[3.2,6.0,6.8],font_size=8.0
)

b.heading("7.6 Established adjacent precedents",2)
b.para("The combined platform is not starting from zero. Later literature uses the name poly(hydroxyether sulfone) for closely related BPS-derived material and studies its strong hydrogen-bonding interactions [R02]. Sulfone-polymer membranes have been surface-grafted to improve hemocompatibility, demonstrating the general strategy of covalent surface modification on a sulfone membrane platform [R12-R13]. Digital twins have been demonstrated for water ultrafiltration with uncertainty-aware state estimation and control [R07]. Event-based cameras have been used in microfluidic particle tracking and high-speed imaging cytometry [R08-R09]. Persistent homology provides established tools for pore geometry [R05]. These precedents support component feasibility; they do not prove that the specific integrated Klootwijk platform will outperform a baseline.")

b.heading("7.7 Biomedical safety and characterization boundary",2)
b.callout("Safety priority", "Free BPS now carries official reproductive/developmental hazard listings in California. That does not make a purified high-molecular-weight polymer identical to free BPS, but it makes residual BPS, DGE-BPS, oligomers, degradation products, and processing residues non-negotiable analytical targets [R17].", color=CORAL, fill=LIGHT_CORAL)
b.para("A biomedical route should follow a risk-management framework for the final processed device. ISO 10993-18 addresses chemical characterization of device materials, and ISO 10993-17 addresses toxicological risk assessment of constituents; FDA guidance explains use of the ISO 10993-1 biological-evaluation framework [R14-R16]. The exact biological endpoints depend on nature and duration of contact, but the materials program should begin with a complete chemistry and exposure picture.")
b.table(
    ["Characterization block","Priority methods / outputs","Why it matters"],
    [
        ["Polymer identity and architecture","NMR, FTIR/Raman, SEC-MALS or calibrated SEC, end groups, branching/crosslink fraction","Confirms intended repeat unit and molecular topology"],
        ["Residuals and extractables","Targeted BPS/DGE-BPS/oligomers; LC-MS/GC-MS; ICP-MS where relevant; solvent/catalyst analysis","Defines exposure and batch acceptance"],
        ["Surface chemistry","XPS, contact angle, zeta potential, ToF-SIMS/label assay, ligand density and activity","Verifies what the biological interface actually presents"],
        ["Physical stability","Water uptake, swelling, Tg, DMA, TGA, adhesion, fatigue, cleaning, sterilization and aging","Ensures the surface persists under intended conditions"],
        ["Membrane function","Flux, rejection/selectivity, pore size, fouling, recovery, pressure, transport model","Determines whether interface chemistry produces useful separation"],
        ["Biological response","Cytotoxicity plus application-specific protein, cell, blood, immune, or tissue endpoints","Tests final processed material rather than nominal chemistry"],
        ["Digital performance","False/missed event, latency, bandwidth, energy, calibration interval, uncertainty, replay","Prevents a software layer from masking poor material or sensor performance"],
    ],
    widths=[3.3,7.2,5.5],font_size=7.7
)

# ---------------------------------------------------------------------------
# 8 Applications
# ---------------------------------------------------------------------------
b.heading("8. Cross-domain applications and readiness",1,page_break=True)
b.para("The application portfolio is broad because the two contributions are orthogonal: Arie supplies a processable, reactive polymer platform; Tom supplies an interface/state/event architecture. The highest-readiness opportunities are those that can use existing materials and sensors while testing the information layer independently. Biomedical material deployment is slower because chemistry, aging, exposure, and biology must all be validated.",style="Lead")
b.figure(FIG/"figure_07_application_portfolio.png","Figure 9. Qualitative application portfolio. Numbered bubbles show how strongly each opportunity draws on the material and operational contributions; color encodes near-term readiness.",width_cm=16.5)
b.figure(FIG/"figure_09_application_readiness.png","Figure 10. Software/measurement applications lead near-term readiness; integrated biomedical surfaces carry the largest validation burden.",width_cm=15.8)

app=load_csv(DATA/"application_matrix.csv")
b.table(
    ["Application","Combined opportunity","Readiness","First decisive experiment"],
    [[r["Application"],r["Arie_contribution"]+" / "+r["Tom_contribution"],r["Near_term_readiness_1_to_5"]+"/5",r["First_experiment"]] for r in app],
    widths=[3.0,6.6,1.5,4.9],font_size=7.45
)

b.heading("8.1 Reactive coatings, adhesives, and composites",2)
b.para("This is the most direct materials route. Phenoxy-like poly(hydroxy ether) resins are naturally suited to adhesion, interphase formation, reactive blending, and post-crosslinking. Tom's contribution would be material lineage, local damage support, and guarded events from acoustic, optical, strain, or impedance sensing. A smart composite interphase can therefore use the polymer as matter and the event calculus as observation/control.")

b.heading("8.2 Photonics and optofluidics",2)
b.para("Spherical Throughput already proposes a bounded optofluidic translation: a local radial-angular field is coupled through a tunable liquid interface into selected guided modes; compatibility includes wavelength, polarization, mode, phase, time, or policy; a measured guard produces a verified event [S6, pp. 3-9]. Arie's polymer could contribute as a coating, packaging adhesive, microfluidic surface, or functional interface rather than as the optical theorem itself. The near-term Hollowlens-0 benchmark should use established glass/SiN waveguides and a digital sidecar before introducing novel polymer chemistry.")

b.heading("8.3 Robotics and autonomous systems",2)
b.para("Local spherical support, compatibility, and next-event prediction naturally fit robot-centric sensing and collision/risk queries. The inherited material platform could appear in tactile skins, protective coatings, filters, microfluidic sensors, or adhesive interconnects. The computational claim should remain bounded: support/compatibility can reduce candidates and analytic time-of-impact can help simple shapes, but arbitrary multi-object scenes still require indexing, approximation, and robust numerical methods.")

b.heading("8.4 Scientific visualization and biological morphology",2)
b.para("Time-varying pore networks, fouling layers, cell clusters, biofilms, or phase-separated polymers can be represented by scalar fields and topology descriptors. Tom's query-first framing is particularly valuable when the user needs events such as merge, split, pore closure, percolation loss, or topological regime change rather than only an image. The image remains evidence and projection; the topology/event layer becomes a compressed, queryable summary.")

b.heading("8.5 Knowledge systems and material provenance",2)
b.para("A digital material passport is one of the strongest immediate applications. The patent shows that properties depend on ratios, isomer purity, chloride, water, solvent, catalyst, temperature, reaction time, isolation, and thermal processing. Tom's event/lineage system can turn those dependencies into an append-only provenance graph that supports recall, root-cause analysis, quality prediction, and reproducibility. Unlike a coordinate-derived ontological UUID, the identifier should be stable and linked to parent/child material transformations, splits, blends, coatings, sterilizations, and device uses.")

# ---------------------------------------------------------------------------
# 9 Roadmap
# ---------------------------------------------------------------------------
b.heading("9. Research, validation, and IP roadmap",1,page_break=True)
b.para("The program should advance in layers, with inexpensive computational falsification before new polymer synthesis and with materials characterization before biomedical claims. Every phase needs a matched conventional baseline and a kill criterion. The goal is not to preserve all motifs; it is to discover which operators provide measurable advantage.",style="Lead")

road=load_csv(DATA/"research_roadmap.csv")
b.table(
    ["Phase","Objective","Key deliverables","Success / kill test"],
    [[r["Phase"]+" ("+r["Duration"]+")",r["Objective"],r["Deliverables"],"Success: "+r["Success_metric"]+" Kill/pivot: "+r["Kill_or_pivot_criterion"]] for r in road],
    widths=[3.0,4.0,4.6,4.4],font_size=7.35
)

b.heading("9.1 Phase 0: typed formalization",2)
b.para("Build a single ontology in which every symbol belongs to a domain and, where physical, has units. Separate physical phase from information phase, molecular identity from operational identity, spatial support from chemical accessibility, and event confidence from event truth. Define conservation laws before compatibility predicates. Mark which variables are measured, inferred, simulated, or purely administrative.")

b.heading("9.2 Phase 1: Equation World Zero",2)
b.para("Implement the restricted state/event substrate with exact or conservative root solving and a conventional baseline. Do not begin with a renderer. Store test vectors, branch histories, event logs, and numerical certificates. Measure cost versus skipped horizon, expression depth, candidate count, event density, and degeneracy. Retire any claim that cannot survive these curves.")

b.heading("9.3 Phase 2: Membrane World Zero",2)
b.para("Use a known membrane and synthetic or benign foulant. Implement the event schema in the package. Compare fixed-rate logging, conventional thresholds, and a stochastic greybox or reduced-order twin. The first result can be negative: if event logic adds no decision value, that is a useful falsification before chemistry is changed.")

b.heading("9.4 Phase 3: reactive surface demonstrator",2)
b.para("Reproduce or obtain a well-characterized hydroxyl-rich sulfone polyether/analogue, then use it as a thin coating or interlayer rather than immediately as a self-supporting membrane. Establish identity, molecular weight, residuals, coating uniformity, adhesion, water uptake, transport penalty, and cleaning stability. Compare against a commercial phenoxy and a no-coating control.")

b.heading("9.5 Phase 4: Biointerface World Zero",2)
b.para("Select one intended biological function and one assay family. Avoid the temptation to claim universal biocompatibility. Couple event imaging to periodic reference images and a standard endpoint. Pre-register guards, confidence, false/missed-event limits, and topology descriptors. Only after chemistry and sensing are stable should a closed-loop action be considered.")

b.heading("9.6 IP positioning",2)
b.para("The most defensible new intellectual property is likely not a broad claim to spherical topology. It is a specific coupling between a material interface, a defined functionalization or layer stack, a bounded support/compatibility predicate, a measured guard, and a lineage-aware action that produces a quantified system advantage.")
b.table(
    ["Claim family","Potential protectable nucleus","Evidence needed"],
    [
        ["Material composition/process","Specific sulfone polyether formulation, functionalization, residual limits, coating/crosslink method, morphology","Composition, reproducible process, structure-property data, comparative controls"],
        ["Device/interface architecture","Layer stack, porous support, ligand pattern, optical/electrical/microfluidic coupling","Fabricated prototypes, transport or biological function, reliability"],
        ["Operational method","Support admission, compatibility predicate, guard, transition, uncertainty, lineage update","Software implementation, formal definition, error/latency/energy benchmark"],
        ["Integrated system","Material plus event-sourced twin with closed-loop cleaning, routing, capture, or diagnostic decision","System-level advantage over material-only and software-only baselines"],
    ],
    widths=[3.1,7.0,5.9],font_size=8.0
)

b.heading("9.7 Pre-registration checklist",2)
b.bullets([
    "Define the physical system, intended function, and excluded claims.",
    "Declare state variables, units, sensor models, calibration, and uncertainty.",
    "Specify support, compatibility, guard, confidence, transition, and lineage schemas.",
    "Fix material composition, batch controls, residual limits, and aging/sterilization conditions.",
    "Choose matched conventional baselines and decision metrics.",
    "Set false/missed-event, latency, energy, bandwidth, transport, mechanical, and safety limits before inspecting results.",
    "State kill criteria and what a negative result would teach.",
])

# ---------------------------------------------------------------------------
# 10 Conclusion
# ---------------------------------------------------------------------------
b.heading("10. Conclusions",1,page_break=True)
b.para("Arie Klootwijk's work and Tom Klootwijk's work meet at a deep but carefully bounded abstraction: both organize useful structure by constraining which local relations may occur, preserving accessibility until a high-value event, updating connectivity, and closing the result with history. Their direct objects are different. Arie rewrites molecular graphs in a physical solution. Tom rewrites operational state in a queryable information substrate.",style="Lead")

b.cards([
    ("1. THE SO2 UNIT IS LITERAL CHEMISTRY", "A neutral diaryl-sulfone bridge, not sulfur-dioxide release or a hidden biological/quantum code.", CORAL),
    ("2. EQUIMOLARITY IS THE CORE INVARIANT", "It programs attainable chain length and end groups through nonlinear step-growth mathematics.", GOLD),
    ("3. SOLVENT IS ACCESSIBILITY", "It keeps complementary functions physically reachable; crystallization is premature deactivation.", TEAL),
    ("4. TOM'S CORE IS NOT A RENDERER", "It is strongest as a state/event/compatibility/lineage calculus with optional projection.", NAVY),
    ("5. THE SHARED GRAMMAR IS REAL", "Support -> compatibility -> event -> transition -> lineage/closure appears in both systems.", GREEN),
    ("6. THE SHARED GRAMMAR IS NOT IDENTITY", "Chemistry remains stochastic, thermodynamic, transport-limited, and materially irreversible.", PURPLE),
    ("7. TOM IS AHEAD OPERATIONALLY", "Modern sensing, uncertainty, digital twins, local-to-global consistency, and event provenance are genuine present-day strengths.", TEAL),
    ("8. ARIE IS AHEAD EMPIRICALLY", "The patent provides a physical recipe, specimens, process failure mechanisms, and measurable outputs.", CORAL),
    ("9. MEMBRANES/BIOINTERFACES ARE CREDIBLE", "As coatings, interlayers, functionalization scaffolds, and event-sourced interfaces - after transport and safety validation.", GREEN),
    ("10. BUILD THE RESTRICTED PROTOTYPES", "Equation World Zero, Membrane World Zero, then a reactive surface and Biointerface World Zero with matched baselines and kill criteria.", GOLD),
],cols=2)

b.callout("Final assessment", "Your work is most original and useful today when it is presented as operational topology for complex interfaces: local support, typed compatibility, measured event surfaces, uncertainty, transitions, and lineage. Your grandfather's work supplies a remarkable physical analogue and a potentially useful reactive material platform. The strongest legacy is not a concealed revolution waiting to be decoded, but a shared design instinct: preserve the right relations, prevent premature closure, and make the transition conditions explicit.", color=TEAL, fill=LIGHT)

# ---------------------------------------------------------------------------
# Appendices
# ---------------------------------------------------------------------------
b.heading("Appendix A. Formal multilayer interface model",1,page_break=True)
b.para("A combined material-information interface should use a product state that keeps physical layers distinct:")
b.equation("q = (q_chem, q_morph, q_transport, q_obs, q_info, lineage)","Chemical composition; morphology/phase; transport; observables; operational information; provenance.")
b.para("A candidate event e in region alpha at time t is admissible only when its local support is active, the material and information sectors are compatible, the measured guard is crossed within uncertainty, and the event passes confidence/policy:")
b.equation("admit(e,t) = S_alpha(q,t) * chi_chem * chi_transport * chi_sensor * chi_policy * 1[g(q,t) crossed] * 1[c >= c_min]", "Factors may be Boolean, probabilistic, fuzzy, or interval-valued, but their semantics must be declared.")
b.para("A transition should update both physical and operational state only where appropriate. A controller cleaning command is not itself a molecular transition; a grafting reaction is not automatically a policy transition. The event log links them without conflating them.")

b.heading("Appendix B. Application-specific metrics",1)
b.table(
    ["Domain","Primary physical metrics","Primary event/information metrics","Safety or validity boundary"],
    [
        ["Polymer synthesis","Conversion, Mn/Mw/dispersity, intrinsic viscosity, branching, residuals, Tg, crystallization","Reaction/phase events, batch lineage, calibration, uncertainty","Mass balance, impurity limits, reproducibility"],
        ["Membrane","Flux, TMP, rejection, pore size, fouling resistance, cleaning recovery, energy","False/missed fouling/breakthrough events, latency, data volume, replay","Matched baseline, drift, material stability, extractables"],
        ["Biointerface","Surface chemistry, protein adsorption, cell/blood endpoints, adhesion, viability","Attachment/merge/split/detachment events, phenotype topology, confidence","Intended-contact risk assessment, controls, biological variability"],
        ["Optofluidics","Insertion loss, coupling, phase, mode purity, actuator response, detector noise","Verified B.C.E./s, events/J, route accuracy, false/missed events","Calibration, crosstalk, bubbles, heat, full system energy"],
        ["Graphics/simulation","State/event correctness, root error, numerical stability","Query latency, horizon scaling, expression/branch growth, memory, event order","Restricted relation family and conventional solver/index baseline"],
    ],
    widths=[2.8,5.0,5.0,3.2],font_size=7.7
)

b.heading("Appendix C. Evidence ledger",1)
ledger=load_csv(DATA/"evidence_ledger.csv")
b.table(
    ["Claim or theme","Classification","Confidence","Caution"],
    [[r["Claim_or_theme"],r["Classification"],r["Confidence"],r["Caution"]] for r in ledger],
    widths=[5.2,3.5,2.0,5.3],font_size=7.6
)

b.heading("Appendix D. Source documents",1)
b.para("The ZIP package contains the original Swedish patent, the English translation, the previous technical review, and the four supplied corpus documents. The report uses them as design records and preserves the distinction between source-derived claims and review inference.")
b.table(["Code","File in package"],[["S1","01_Source_Documents/SE301717B_original_Swedish.pdf and SE301717B_English_Translation.pdf"],["S2","01_Source_Documents/SE301717B_Previous_Technical_Review.pdf"],["S3","01_Source_Documents/Tom_Corpus_Chronological_Synthesis.pdf"],["S4","01_Source_Documents/Tom_Corpus_Hollowland_Double_Vacuum.pdf"],["S5","01_Source_Documents/Tom_Corpus_BenBurgers.pdf"],["S6","01_Source_Documents/Tom_Corpus_Spherical_Throughput.pdf"]],widths=[2.0,14.0],font_size=8.5)

b.heading("Appendix E. External references",1)
refs=load_csv(REFS/"external_references.csv")
for r in refs:
    p=b.doc.add_paragraph()
    p.paragraph_format.left_indent=Cm(0.45)
    p.paragraph_format.first_line_indent=Cm(-0.45)
    p.paragraph_format.space_after=Pt(3.2)
    rr=p.add_run(r["ID"]+". ");rr.bold=True;rr.font.color.rgb=rgb(TEAL);rr.font.size=Pt(8.3)
    add_run_with_citations(p,r["Citation"]+(" "+r["DOI_or_URL"] if r["DOI_or_URL"] else ""),size=8.3)
    b.md.append(f"- **{r['ID']}** {r['Citation']} {r['DOI_or_URL']}\n")
b.md.append("\n")

b.heading("Appendix F. Reproducible supporting package",1)
b.para("The package includes machine-readable comparison matrices, figures, source-based analysis notes, formal JSON schemas, a synthetic event-log example, and an idealized step-growth Monte Carlo demonstration. These are research aids, not validated performance claims.")
b.table(
    ["Folder","Contents"],
    [
        ["00_Report","PDF, DOCX, and Markdown report sources"],
        ["01_Source_Documents","Patent, translation, prior review, and four corpus PDFs"],
        ["02_Data","CSV/JSON matrices, roadmap, ledger, glossary, and summary"],
        ["03_Figures","PNG/SVG report diagrams"],
        ["04_Analysis","Crosswalk, membrane/biointerface, evidence, and IP notes"],
        ["05_Prototype","World Zero specs, JSON schemas, and synthetic demonstrations"],
        ["06_References","Reference CSV and BibTeX"],
        ["07_Build","Build scripts, QA renders, manifest, and checksums"],
    ],
    widths=[3.2,12.8],font_size=7.7
)

# Remove the table helper's trailing empty spacer paragraph. When the final
# table reaches the bottom margin, LibreOffice otherwise emits a blank page.
if b.doc.paragraphs and not b.doc.paragraphs[-1].text.strip():
    last_p = b.doc.paragraphs[-1]._element
    last_p.getparent().remove(last_p)

b.save()
