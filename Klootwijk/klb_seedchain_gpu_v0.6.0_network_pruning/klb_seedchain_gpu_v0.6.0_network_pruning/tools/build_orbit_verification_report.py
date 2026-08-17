from __future__ import annotations

import csv
import hashlib
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "KLB_OrbitSeed_GPU_v0.3.0_Laptop_Verification_ELI5.docx"
SUMMARY_CSV = ROOT / "orbit_verification_metrics.csv"

# standard_business_brief preset with an editorial-cover title override.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "1A1A1A"
MUTED = "5E6670"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
POSITIVE = "1F3A5F"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def load_modes(preset: str):
    path = ROOT / f"orbit_{preset}_results.csv"
    with path.open(newline="", encoding="utf-8") as handle:
        return {row["mode"]: row for row in csv.DictReader(handle)}


def load_telemetry(name: str):
    path = ROOT / f"orbit_{name}_gpu_telemetry.csv"
    rows = []
    with path.open(newline="", encoding="utf-8") as handle:
        for row in csv.reader(handle):
            if len(row) < 8:
                continue
            rows.append(
                {
                    "util": float(row[2].strip()),
                    "mem": float(row[3].strip()),
                    "power": float(row[4].strip()),
                    "temp": float(row[5].strip()),
                    "sm": float(row[6].strip()),
                    "memclk": float(row[7].strip()),
                }
            )
    active = [row for row in rows if row["util"] >= 50.0]
    return {
        "samples": len(rows),
        "active_samples": len(active),
        "util_avg": sum(r["util"] for r in active) / len(active),
        "power_avg": sum(r["power"] for r in active) / len(active),
        "power_max": max(r["power"] for r in rows),
        "temp_min": min(r["temp"] for r in rows),
        "temp_max": max(r["temp"] for r in rows),
        "mem_max": max(r["mem"] for r in rows),
        "sm_avg": sum(r["sm"] for r in active) / len(active),
        "sm_max": max(r["sm"] for r in rows),
        "memclk_max": max(r["memclk"] for r in rows),
    }


def f(row, key):
    return float(row[key])


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color, bold=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + side
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DBE2", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    assert sum(widths_dxa) == CONTENT_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440.0)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.0, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths_dxa, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_GRAY)
        align = WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=NAVY, size=9.0, align=align)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=8.85, align=align)
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    return table


def add_numbering_definition(doc, *, bullet):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "\u2022" if bullet else "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, *, bold_prefix=None, italic=False, color=INK, after=6, align=None):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True, color=color)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, italic=italic, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, italic=italic, color=color)
    return paragraph


def add_callout(doc, label, text, *, fill=CALLOUT, accent=POSITIVE):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.10
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement("w:" + edge)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:space"), "6")
        node.set(qn("w:color"), "CDD5DF")
        borders.append(node)
    p_pr.append(borders)
    label_run = paragraph.add_run(label + "  ")
    set_run_font(label_run, size=11, color=accent, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=11, color=INK)
    return paragraph


def add_field(paragraph, instruction):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    rpr.append(color)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    rpr.append(size)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def set_body_footer(section):
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    set_run_font(run, size=9, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def configure_section(section, top=1.0):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(top)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_document(doc):
    section = doc.sections[0]
    configure_section(section, top=1.0)
    section.different_first_page_header_footer = True
    footer = section.first_page_footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Independent laptop verification | 16 August 2026")
    set_run_font(run, size=9, color=MUTED)

    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    # Keep automatic field updates disabled at open time. Word otherwise shows
    # an interactive security/update prompt in automation mode. The PDF export
    # pass updates the local PAGE/NUMPAGES footer fields explicitly.
    update.set(qn("w:val"), "false")


def start_body_page(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, top=0.55)
    section.different_first_page_header_footer = False
    set_body_footer(section)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(11)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run("KLB OrbitSeed GPU 0.3.0  |  Independent laptop verification")
    set_run_font(run, size=9, color=MUTED, bold=True)
    return section


def write_summary(file_modes, laptop_modes, vram_modes, telemetry):
    rows = []
    for preset, modes in (("file", file_modes), ("laptop", laptop_modes), ("vram", vram_modes)):
        seed = modes["query_seed_direct"]
        end = modes["materialize_plus_query"]
        dense = modes["query_dense"]
        compact = modes["compact_seed_events"]
        rows.extend(
            [
                (preset, "dense_working_set_mib", f(seed, "dense_bytes") / 1048576.0, "MiB"),
                (preset, "working_set_ratio", f(seed, "working_set_ratio"), "x"),
                (preset, "direct_seed_p50", f(seed, "p50_ms"), "ms"),
                (preset, "end_to_end_dense_p50", f(end, "p50_ms"), "ms"),
                (preset, "direct_p50_advantage", 100.0 * (1.0 - f(seed, "p50_ms") / f(end, "p50_ms")), "%"),
                (preset, "direct_vs_resident_dense", f(seed, "p50_ms") / f(dense, "p50_ms"), "x slower"),
                (preset, "compacted_events", int(compact["event_count"]), "events"),
            ]
        )
    for name, stats in telemetry.items():
        rows.extend(
            [
                (name, "active_power_average", stats["power_avg"], "W"),
                (name, "power_maximum", stats["power_max"], "W"),
                (name, "temperature_maximum", stats["temp_max"], "C"),
                (name, "vram_used_maximum", stats["mem_max"], "MiB"),
                (name, "active_sm_clock_average", stats["sm_avg"], "MHz"),
            ]
        )
    rows.extend(
        [
            ("validation", "manifest_hashes", 82, "of 82 passed"),
            ("validation", "cpu_test_suites", 2, "of 2 passed"),
            ("validation", "compute_sanitizer_errors", 0, "errors"),
            ("validation", "full_horizon_support_difference", 1, "candidate"),
            ("profile", "seed_count_registers", 46, "registers/thread"),
            ("profile", "seed_compact_registers", 52, "registers/thread"),
            ("profile", "kernel_spills", 0, "bytes"),
            ("profile", "seed_count_achieved_occupancy", 81.636171, "%"),
            ("profile", "seed_compact_achieved_occupancy", 65.610496, "%"),
            ("profile", "seed_count_compute_throughput", 66.187378, "% of peak"),
            ("profile", "seed_count_dram_throughput", 0.002975, "% of peak"),
            ("profile", "seed_count_l2_throughput", 0.095049, "% of peak"),
            ("profile", "seed_count_xu_pipe_utilization", 20.064937, "% of peak active"),
            ("profile", "seed_count_idc_request_activity", 63.359952, "% of peak elapsed"),
            ("profile", "seed_compact_compute_throughput", 65.536077, "% of peak"),
            ("profile", "seed_compact_dram_throughput", 0.007861, "% of peak"),
            ("profile", "seed_compact_l2_throughput", 0.081524, "% of peak"),
            ("profile", "seed_compact_xu_pipe_utilization", 18.105335, "% of peak active"),
            ("profile", "compact_global_atomic_requests", 1243, "requests"),
            ("profile", "compact_l2_atomic_input_activity", 0.038542, "% of peak elapsed"),
        ]
    )
    with SUMMARY_CSV.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(("scope", "metric", "value", "unit_or_result"))
        writer.writerows(rows)


def build_report():
    file_modes = load_modes("file")
    laptop_modes = load_modes("laptop")
    vram_modes = load_modes("vram")
    telemetry = {
        "demo": load_telemetry("demo"),
        "laptop": load_telemetry("laptop"),
        "vram": load_telemetry("vram"),
    }
    write_summary(file_modes, laptop_modes, vram_modes, telemetry)

    doc = Document()
    configure_document(doc)
    bullet_id = add_numbering_definition(doc, bullet=True)
    number_id = add_numbering_definition(doc, bullet=False)
    props = doc.core_properties
    props.title = "KLB OrbitSeed GPU 0.3.0 Laptop Verification"
    props.subject = "ELI5 CUDA, compression, crossover, correctness, and feasibility report"
    props.author = "OpenAI Codex"
    props.keywords = "KLB, OrbitSeed, CUDA, RTX 5070 Ti, orbit, verification, ELI5"

    # Cover page.
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(17)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("INDEPENDENT LAPTOP VERIFICATION")
    set_run_font(run, size=10, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("KLB OrbitSeed GPU 0.3.0")
    set_run_font(run, size=29, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(7)
    run = subtitle.add_run("Does compact orbit reconstruction beat a dense GPU timeline?")
    set_run_font(run, size=14, color=DARK_BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(22)
    run = meta.add_run("Human-readable / ELI5 edition | RTX 5070 Ti Laptop | 16 August 2026")
    set_run_font(run, size=10, color=MUTED, italic=True)
    add_callout(
        doc,
        "BOTTOM LINE",
        "The required first-run checks pass and the direct seed path narrowly beats rebuilding and then querying a dense timeline: 0.93% at 295 MiB, 2.19% at 512 MiB, and 2.50% at 2 GiB. It does not beat a dense table that is already resident, and it is not yet an operational orbit product: a stronger full-horizon oracle finds one CPU/GPU support mismatch and no SGP4 accuracy comparison exists.",
        fill=LIGHT_BLUE,
    )
    metrics = doc.add_table(rows=2, cols=3)
    set_table_geometry(metrics, [3120, 3120, 3120])
    set_repeat_table_header(metrics.rows[0])
    for idx, (label, value) in enumerate(
        (("512 MiB CROSSOVER", "+2.19%"), ("EVENT SET", "1,243 exact"), ("SANITIZER", "0 errors"))
    ):
        set_cell_shading(metrics.cell(0, idx), LIGHT_GRAY)
        set_cell_text(metrics.cell(0, idx), label, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(metrics.cell(1, idx), value, bold=True, color=POSITIVE, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body(
        doc,
        "ELI5 analogy: carrying a tiny recipe and cooking the answer on demand just beats unpacking a warehouse-sized freezer for one meal. If the meal is already on the table, the recipe is slower.",
        italic=True,
        color=MUTED,
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    # Page 2.
    start_body_page(doc)
    add_heading(doc, "1. Verdict in plain English", 1)
    add_callout(
        doc,
        "VERDICT",
        "PASS as a CUDA compression/crossover experiment. CONDITIONAL as a coarse scheduling engine. NO-GO for navigation, collision avoidance, precise antenna pointing, or any claim of verified SGP4-quality ephemerides.",
    )
    add_heading(doc, "What it actually does", 2)
    add_body(doc, "The file stores 32 GPS orbital mean-element seeds plus seven daily chain nodes in 3,809 bytes. At query time, the GPU reconstructs two adjacent satellite positions, checks station range and route compatibility, evaluates a 10-degree elevation guard, and emits acquisition/loss events.")
    add_body(doc, "The seven-day test spans 604,800 one-second intervals and 19,353,600 satellite-time candidates. A dense float4 position timeline for the same declared horizon is 309,658,112 bytes (295.3 MiB). The larger stress presets repeat that bounded timeline only to create load.")
    add_heading(doc, "The four acceptance gates", 2)
    acceptance = (
        ("CPU/GPU oracle prefix", "PASS", "4,096 epochs; exact aggregate counters"),
        ("Direct/dense counters", "PASS", "Equal in file, 512 MiB, and 2 GiB runs"),
        ("Compacted event sets", "PASS", "717 / 1,243 / 4,970 sorted events matched"),
        ("Event output truncated", "PASS", "No, for every measured run"),
    )
    add_table(doc, ("Gate", "Result", "Evidence"), acceptance, [3000, 1200, 5160], numeric_cols=(1,))
    add_heading(doc, "What the result does not mean", 2)
    for text in (
        "It is model-based reconstruction, not lossless compression of a pre-existing 295 MiB trajectory file.",
        "The current predictor is fixed-iteration Kepler plus secular J2, not SGP4 and not navigation-grade.",
        "An 81,296x horizon ratio does not mean the predictor is free; it trades stored positions for trigonometric GPU work.",
    ):
        add_list_item(doc, text, bullet_id)

    # Page 3.
    start_body_page(doc)
    add_heading(doc, "2. Test platform and build evidence", 1)
    add_table(
        doc,
        ("Item", "Measured configuration"),
        (
            ("GPU", "NVIDIA GeForce RTX 5070 Ti Laptop GPU; compute capability 12.0"),
            ("Memory / topology", "12,227 MiB reported global memory; 46 SMs; 36 MiB L2; 192-bit bus"),
            ("Driver", "NVIDIA 591.59"),
            ("CUDA compiler", "12.8.61 explicitly selected"),
            ("Host compiler", "MSVC 19.44 / Visual Studio 2022 Build Tools"),
            ("Code generation", "Native sm_120 cubin plus compute_120 PTX"),
        ),
        [2200, 7160],
    )
    add_heading(doc, "Build outcome", 2)
    add_body(doc, "The package's build_windows.ps1 was first run unchanged. CMake reported that no CUDA compiler was found, MSBuild then reported a missing project, CTest found no tests, and the helper still exited 0. That is false success and must be fixed.")
    add_callout(doc, "RECOVERY USED", "A separate Visual Studio build directory explicitly selected the installed CUDA 12.8 toolset. No source or algorithm change was needed. That build succeeded, both CPU test suites passed, and the executable contains both sm_120 machine code and compute_120 PTX.", fill=LIGHT_BLUE)
    add_heading(doc, "Static compiler evidence", 2)
    add_table(
        doc,
        ("Kernel", "Registers/thread", "Spills", "Shared bytes"),
        (
            ("Materialize dense", "32", "0", "0"),
            ("Direct seed count", "46", "0", "1,152"),
            ("Resident dense count", "33", "0", "1,152"),
            ("Direct seed compact", "52", "0", "1,152"),
            ("Resident dense compact", "40", "0", "1,152"),
        ),
        [3800, 1960, 1400, 2200],
        numeric_cols=(1, 2, 3),
    )
    add_body(doc, "The direct seed kernels use more registers because they reconstruct orbital state and evaluate trigonometric guards. Register spills were zero. Nsight later measured 81.64% achieved occupancy for direct counting and 65.61% for direct compact output.", italic=True, color=MUTED)

    # Page 4.
    start_body_page(doc)
    add_heading(doc, "3. Actual seven-day application", 1)
    add_body(doc, "The CPU application rebuilt a 52 N, 5 E coarse acquisition/loss schedule over the complete seven-day horizon. It processed 19,353,600 candidate intervals and emitted 717 events. The GPU benchmark then measured six paths over the same file horizon.")
    file_rows = []
    for mode, label in (
        ("query_seed_direct", "Direct seed query"),
        ("materialize_dense", "Materialize dense"),
        ("query_dense", "Query resident dense"),
        ("materialize_plus_query", "Materialize + query"),
        ("compact_seed_events", "Direct + compact"),
        ("compact_dense_events", "Dense + compact"),
    ):
        row = file_modes[mode]
        file_rows.append((label, f"{f(row, 'p50_ms'):.3f}", f"{f(row, 'p95_ms'):.3f}", f"{f(row, 'p99_ms'):.3f}", f"{f(row, 'mean_ms'):.3f}"))
    add_table(doc, ("Mode", "p50 ms", "p95 ms", "p99 ms", "Mean ms"), file_rows, [3300, 1515, 1515, 1515, 1515], numeric_cols=(1, 2, 3, 4))
    seed = file_modes["query_seed_direct"]
    end = file_modes["materialize_plus_query"]
    dense = file_modes["query_dense"]
    add_callout(doc, "CROSSOVER", f"Direct seed was {f(seed, 'p50_ms'):.3f} ms versus {f(end, 'p50_ms'):.3f} ms end-to-end dense: a {100*(1-f(seed,'p50_ms')/f(end,'p50_ms')):.2f}% p50 advantage while avoiding a 295.3 MiB table. But it was {f(seed,'p50_ms')/f(dense,'p50_ms'):.2f}x slower than querying a dense table that already existed.", fill=LIGHT_BLUE)
    add_heading(doc, "Useful CPU output", 2)
    add_table(
        doc,
        ("Counter", "Result"),
        (
            ("Candidate intervals", "19,353,600"),
            ("Support / compatible intervals", "19,214,155 / 19,214,155"),
            ("Visible sample states", "5,498,030"),
            ("Acquisition/loss events", "717"),
        ),
        [5200, 4160],
        numeric_cols=(1,),
    )

    # Page 5.
    start_body_page(doc)
    add_heading(doc, "4. Sustained 512 MiB laptop result", 1)
    add_body(doc, "The required laptop preset processed 33,554,432 satellite-time candidates and held a 512.000 MiB dense baseline. Each latency sample auto-repeated until it reached at least 250 ms; eleven samples were recorded.")
    laptop_rows = []
    for mode, label in (
        ("query_seed_direct", "Direct seed query"),
        ("materialize_dense", "Materialize dense"),
        ("query_dense", "Query resident dense"),
        ("materialize_plus_query", "Materialize + query"),
        ("compact_seed_events", "Direct + compact"),
        ("compact_dense_events", "Dense + compact"),
    ):
        row = laptop_modes[mode]
        laptop_rows.append((label, f"{f(row,'p50_ms'):.3f}", f"{f(row,'p95_ms'):.3f}", f"{f(row,'p99_ms'):.3f}", row["inner_repeats"]))
    add_table(doc, ("Mode", "p50 ms", "p95 ms", "p99 ms", "Inner repeats"), laptop_rows, [3300, 1450, 1450, 1450, 1710], numeric_cols=(1, 2, 3, 4))
    seed = laptop_modes["query_seed_direct"]
    end = laptop_modes["materialize_plus_query"]
    dense = laptop_modes["query_dense"]
    compact = laptop_modes["compact_seed_events"]
    add_callout(doc, "LAPTOP DECISION", f"Direct seed was {f(seed,'p50_ms'):.3f} ms and end-to-end dense was {f(end,'p50_ms'):.3f} ms: direct won by {100*(1-f(seed,'p50_ms')/f(end,'p50_ms')):.2f}%. It delivered {f(seed,'candidates_per_second')/1e9:.3f} billion candidates/s. The compact path emitted {int(compact['event_count']):,} exact events without truncation.", fill=LIGHT_BLUE)
    add_heading(doc, "Why this is only a narrow win", 2)
    add_body(doc, f"A resident dense query took only {f(dense,'p50_ms'):.3f} ms, making direct reconstruction {f(seed,'p50_ms')/f(dense,'p50_ms'):.2f}x slower for repeated queries against the same table. Direct reconstruction is useful when materialization is avoided or the dense working set cannot remain resident; it is not a blanket replacement for reuse.")
    add_body(doc, f"Warp-compacted direct output cost {f(compact,'p50_ms'):.3f} ms, {100*(f(compact,'p50_ms')/f(seed,'p50_ms')-1):.2f}% above counting alone even though event yield was only {100*f(compact,'event_yield'):.6f}%. Compaction overhead is measurable despite sparse output.")

    # Page 6.
    start_body_page(doc)
    add_heading(doc, "5. Optional 2 GiB comparison and scaling", 1)
    add_body(doc, "The optional vram preset was safe to run with more than 11.5 GiB initially free. It processed 134,217,728 candidates and allocated a 2,048.000 MiB dense position baseline. The bounded seven-day timeline was repeated for load; this is not a 48-day physical prediction.")
    compare_rows = []
    for preset, label, modes in (
        ("file", "Actual 7-day file", file_modes),
        ("laptop", "512 MiB laptop", laptop_modes),
        ("vram", "2 GiB vram", vram_modes),
    ):
        seed = modes["query_seed_direct"]
        end = modes["materialize_plus_query"]
        dense = modes["query_dense"]
        compact = modes["compact_seed_events"]
        compare_rows.append(
            (
                label,
                f"{f(seed,'dense_bytes')/1048576:.1f}",
                f"{f(seed,'p50_ms'):.3f}",
                f"{f(end,'p50_ms'):.3f}",
                f"{100*(1-f(seed,'p50_ms')/f(end,'p50_ms')):.2f}%",
                f"{f(seed,'p50_ms')/f(dense,'p50_ms'):.2f}x",
                f"{int(compact['event_count']):,}",
            )
        )
    add_table(doc, ("Preset", "Dense MiB", "Direct", "End dense", "Direct win", "vs resident", "Events"), compare_rows, [2300, 1100, 1250, 1400, 1300, 1200, 810], numeric_cols=(1, 2, 3, 4, 5, 6))
    add_callout(doc, "SCALING RESULT", "Direct reconstruction stayed near 5.06-5.25 billion candidate intervals/s. Its p50 advantage over materialize-plus-query grew from 0.93% to 2.50%, but remained small. The experiment demonstrates a crossover, not a landslide.", fill=LIGHT_BLUE)
    add_heading(doc, "ELI5: one question versus many", 2)
    add_body(doc, "Direct seed is like reading a compact recipe and cooking one answer. Materialize-plus-query is like unpacking an entire freezer before checking one item. For a one-off question, the recipe narrowly wins and needs almost no timeline storage.")
    add_body(doc, "A resident dense table is like food already laid out. Repeated questions are then much cheaper: the dense query is roughly 4.5x faster. The practical crossover therefore depends on reuse, memory pressure, transfer cost, query count, and required accuracy - not on compression ratio alone.")
    add_heading(doc, "Logical bandwidth versus hardware traffic", 2)
    add_body(doc, "The benchmark's arithmetic reports about 100.6-101.1 GB/s logical materialization traffic and 728-788 GB/s logical dense-query traffic. Those are workload byte counts divided by normal-run time. The later seed-kernel Nsight profile measured only 0.0030% peak DRAM throughput and 0.095% peak L2 throughput, confirming that direct reconstruction is compute-heavy rather than external-memory-heavy.", italic=True, color=MUTED)

    # Page 7.
    start_body_page(doc)
    add_heading(doc, "6. Correctness audit: required pass, stronger check fails", 1)
    checks = (
        "All 82 package-manifest SHA-256 entries matched.",
        "KLOC1 payload and seven-node hash chain validated.",
        "Both CPU test suites passed (2 of 2).",
        "Required 4,096-epoch CPU/GPU oracle prefix matched exactly in all three primary runs.",
        "Direct and dense GPU counters matched for file, laptop, and vram presets.",
        "Sorted compact event payloads matched: 717, 1,243, and 4,970 events.",
        "No compact event output was truncated.",
        "Compute Sanitizer memcheck reported zero errors on all orbit modes.",
    )
    for item in checks:
        add_list_item(doc, item, bullet_id)
    add_callout(
        doc,
        "CRITICAL FINDING",
        "Extending --verify-epochs from the required 4,096 prefix to all 604,800 intervals failed. CPU support/compatibility counters were 19,214,155; GPU counters were 19,214,154. Verified event counts remained 717 on both sides. Binary search found the last passing prefix at 435,260 epochs and the first failing prefix at 435,261, around 2026-08-21 06:27:32.693 UTC.",
        fill="FDECEC",
        accent=RED,
    )
    add_heading(doc, "Cross-compiler CPU schedule", 2)
    add_body(doc, "The regenerated MSVC CPU pass CSV retained the same 718 lines, event identities, order, and count as the supplied GCC-generated file, but it was not byte-identical. Three fields changed: one interpolated event time by 0.007813 seconds and two minimum guard values by 0.000001.")
    add_body(doc, "This looks like a floating-point boundary/format determinism issue rather than an event-set failure, but the exact root candidate has not been isolated. The package should not claim complete cross-compiler byte identity or full-horizon exact CPU/GPU counters until that boundary is fixed or guarded by a declared tolerance.", italic=True, color=MUTED)

    # Page 8.
    start_body_page(doc)
    add_heading(doc, "7. Profiling, power, and thermals", 1)
    add_callout(doc, "NSIGHT RESULT", "PASS after the NVIDIA counter permission change and Windows reboot. The prescribed full profile produced a valid 10,043,271-byte orbit_seed_profile.ncu-rep. A focused compact-kernel profile produced a second valid report so event-append contention could be measured instead of inferred.", fill=LIGHT_BLUE)
    add_table(
        doc,
        ("Hardware counter", "Seed count", "Seed compact", "Interpretation"),
        (
            ("Registers / thread", "46", "52", "Compaction costs six registers"),
            ("Achieved occupancy", "81.64%", "65.61%", "Register limit reduces resident warps"),
            ("Compute throughput", "66.19%", "65.54%", "Compute-heavy kernel"),
            ("DRAM throughput", "0.0030%", "0.0079%", "Not DRAM-bandwidth-bound"),
            ("L2 throughput", "0.095%", "0.082%", "Very light L2 pressure"),
            ("XU / special-function pipe", "20.06%", "18.11%", "Transcendental work is material"),
            ("Global atomic requests", "0", "1,243", "Exactly one per emitted event"),
            ("L2 atomic-input activity", "0.043%", "0.039%", "No measurable append bottleneck"),
        ),
        [2600, 1450, 1550, 3760],
        numeric_cols=(1, 2),
    )
    add_body(doc, "Profiler replay ran near 947 MHz and inflated the benchmark-displayed profiled time; the normal sustained timing tables remain the latency source of record. For the count kernel, 41.57% of scheduler cycles had no eligible warp. The largest average stall indicators per issued instruction were branch resolving (3.05 warps), MIO throttle (2.55), and short scoreboard (2.49); long-scoreboard and global-load throttle were both zero.", italic=True, color=MUTED)
    add_body(doc, "Nsight Compute did not expose a direct constant-cache hit-rate metric for this CC 12.0 report. It did expose 63.36% IDC request activity and 0.45% uniform-pipe utilization, while cubin inspection confirmed the 32 KiB constant allocation. Those figures are not a cache hit rate, so constant-cache efficiency remains unquantified rather than guessed.", italic=True, color=MUTED)
    telemetry_rows = []
    for name, label in (("demo", "7-day file"), ("laptop", "512 MiB"), ("vram", "2 GiB")):
        stats = telemetry[name]
        telemetry_rows.append(
            (
                label,
                f"{stats['util_avg']:.0f}%",
                f"{stats['power_avg']:.2f}",
                f"{stats['power_max']:.2f}",
                f"{stats['temp_max']:.0f}",
                f"{stats['mem_max']:.0f}",
                f"{stats['sm_avg']:.0f}",
            )
        )
    add_table(doc, ("Run", "Active GPU", "Avg W", "Max W", "Max C", "Max MiB", "Avg SM MHz"), telemetry_rows, [1900, 1150, 1200, 1200, 1050, 1350, 1510], numeric_cols=(1, 2, 3, 4, 5, 6))
    add_body(doc, "Telemetry was sampled with nvidia-smi every 200 ms and summarizes samples at or above 50% GPU utilization. It is useful whole-device context, not per-kernel energy metering. The 2 GiB run reached 84 C; average active SM clock remained about 2.59 GHz, so no clock collapse was observed, but thermal headroom on the laptop is limited.", italic=True, color=MUTED)
    add_heading(doc, "Static evidence and measured compaction cost", 2)
    for text in (
        "The direct count kernel uses 46 registers/thread; compact direct uses 52.",
        "The resident dense count kernel uses 33 registers/thread; dense compact uses 40.",
        "All five orbit kernels report zero spill stores and loads.",
        "The seed and node arrays are compiled into CUDA constant memory; this Nsight/CC 12.0 metric set did not provide a direct constant-cache hit rate.",
        "The compact path adds roughly 13-16% over direct counting and 64-67% over resident dense counting in these sparse-event tests, but the atomic append itself is not the limiter at 0.003704% yield.",
    ):
        add_list_item(doc, text, bullet_id)

    # Page 9.
    start_body_page(doc)
    add_heading(doc, "8. Compression and feasibility boundary", 1)
    add_heading(doc, "Why the giant ratio is mathematically real", 2)
    add_body(doc, "The 3,809-byte KLOC1 file contains a closed model: 32 fixed-size orbital seeds and seven timeline nodes. Repeating the model at one-second intervals produces 604,801 state samples per satellite. Comparing that tiny model with a chosen dense float4 expansion yields 81,296x for seven days.")
    add_heading(doc, "Why it is not conventional compression", 2)
    add_body(doc, "No 295 MiB measured trajectory was fed into an encoder and recovered losslessly. The source OMM CSV itself is only 4,852 bytes. The system is closer to storing a physics recipe than squeezing an existing movie. Increasing the requested horizon increases the apparent ratio even though the model file barely changes.")
    add_callout(doc, "FAIR CLAIM", "OrbitSeed can avoid materializing hundreds of MiB or GiB of positions for one-off coarse event queries. It cannot yet claim operational ephemeris compression, because the predictor's coordinate and event-order error versus SGP4 has not been measured.", fill=LIGHT_BLUE)
    add_heading(doc, "Accuracy boundary", 2)
    add_body(doc, "CelesTrak general-perturbations mean elements are associated with SGP4. This package instead uses a deterministic fixed-iteration Kepler solve plus precomputed secular J2 rates. Its generated pass CSV is therefore a coarse workload result, not an operational satellite schedule.")
    add_body(doc, "Before deployment, compare coordinates, elevation guards, acquisition/loss identity, and event ordering against an SGP4 reference. A packed predictor is acceptable only when its maximum coordinate and guard error stays below a declared application margin.")
    add_heading(doc, "Feasibility decision", 2)
    add_table(
        doc,
        ("Use", "Decision", "Reason"),
        (
            ("GPU crossover research", "GO", "Correct event equivalence; direct path narrowly beats end-to-end dense"),
            ("Memory-limited coarse preview", "CONDITIONAL", "Useful if one-off queries and approximate orbit model are acceptable"),
            ("Repeated queries on one timeline", "DENSE MAY WIN", "Resident dense query is about 4.5x faster"),
            ("Operational orbital decisions", "NO-GO", "No SGP4 error budget; full-horizon exactness issue remains"),
        ),
        [2500, 1800, 5060],
    )

    # Page 10.
    start_body_page(doc)
    add_heading(doc, "9. Practical uses and poor fits", 1)
    add_heading(doc, "Good experimental matches", 2)
    good = (
        ("Interactive ground-station what-if tools", "Change station, elevation mask, or route and answer occasional visibility questions without allocating a complete timeline."),
        ("GPU event-pipeline research", "Study direct reconstruction, guard evaluation, lineage, warp compaction, and sparse append behavior on a real structured dataset."),
        ("Memory-limited constellation dashboards", "Keep compact seeds resident and reconstruct only requested slices, provided approximate propagation is acceptable."),
        ("Training, replay, and synthetic mission rehearsal", "Use deterministic, repeatable coarse motion where fidelity is explicitly bounded and nobody mistakes it for navigation data."),
        ("Future GPU SGP4 substrate", "Retain KLOC1 identity, query ABI, lineage, and support/guard stages while replacing the predictor with validated SGP4 coefficients or kernels."),
    )
    for title_text, detail in good:
        add_heading(doc, title_text, 3)
        add_body(doc, detail)
    add_heading(doc, "Poor matches", 2)
    for text in (
        "Navigation, collision avoidance, safety of flight, precise antenna steering, or authoritative operations.",
        "Workloads that repeatedly query the same interval and can keep a dense table resident.",
        "Scenarios where external updates are frequent enough that compact seeds no longer describe the state between refreshes.",
        "Claims of lossless trajectory compression when only a compact generative model was stored.",
        "Cross-platform deterministic archives until floating-point and build-tool behavior are canonicalized.",
    ):
        add_list_item(doc, text, bullet_id)

    # Page 11.
    start_body_page(doc)
    add_heading(doc, "10. Recommended next engineering steps", 1)
    recommendations = (
        "Fix build_windows.ps1 so every non-zero native command terminates the script and add an explicit CUDA-toolkit selector.",
        "Make the full 604,800-epoch CPU/GPU oracle exact, or define and enforce a conservative support/guard tolerance tied to the event margin.",
        "Canonicalize cross-compiler floating-point evaluation and CSV formatting if byte-identical regeneration remains a requirement.",
        "Reduce compact-kernel register pressure or test a split compaction design; its 52 registers/thread lower achieved occupancy from 81.64% to 65.61%.",
        "If constant-cache hit rate is decision-critical, map a supported CC 12.0 counter or run a controlled constant-versus-global-memory A/B benchmark; the full Nsight set did not expose a direct hit-rate metric.",
        "Build an SGP4 CPU reference comparison for all source objects across the seven-day horizon, including event identity/order and maximum elevation-guard error.",
        "Profile a GPU SGP4 coefficient path or homogeneous near-Earth/deep-space queues, then repeat the direct-versus-dense crossover at equal accuracy.",
        "Measure multi-station and repeated-query workloads so the materialization/reuse crossover is expressed as a query-count rule, not one timing ratio.",
        "Add controlled energy sampling or vendor telemetry integration if joules per candidate matters; 200 ms nvidia-smi samples are too coarse for per-kernel attribution.",
    )
    for item in recommendations:
        add_list_item(doc, item, number_id)
    add_callout(doc, "GO / NO-GO RULE", "Proceed when direct reconstruction is faster than materialize-plus-query at equal validated accuracy, the dense table cannot be profitably reused, and coordinate/guard error remains below the application margin. Otherwise keep the dense representation or use a validated SGP4 path.", fill="FFF8E8", accent=GOLD)

    # Page 12.
    start_body_page(doc)
    add_heading(doc, "Appendix A. Evidence and deliverables", 1)
    add_body(doc, "Primary result files created in the package root:")
    for text in (
        "orbit_file_results.csv - actual included seven-day GPU run.",
        "orbit_laptop_results.csv - required sustained 512 MiB comparison.",
        "orbit_vram_results.csv - optional 2 GiB comparison.",
        "demo_orbit_console.txt - complete seven-day application console output.",
        "stress_orbit_laptop_console.txt - complete laptop-preset console output.",
        "stress_orbit_vram_console.txt - complete vram-preset console output.",
        "orbit_*_gpu_telemetry.csv - 200 ms nvidia-smi samples for all primary runs.",
        "profile_orbit_console.txt, profile_orbit_rerun_console.txt, and profile_orbit_third_attempt_console.txt - three pre-reboot Nsight permission failures.",
        "profile_orbit_post_reboot_console.txt and orbit_seed_profile.ncu-rep - successful prescribed full seed profile.",
        "profile_orbit_compact_console.txt and orbit_seed_compact_profile.ncu-rep - successful focused compact-event profile.",
        "orbit_seed_ncu_raw.csv and orbit_seed_compact_ncu_raw.csv - raw exported hardware counters.",
        "orbit_seed_profile_details.txt and orbit_seed_compact_profile_details.txt - human-readable counter summaries.",
        "orbit_cuobjdump_resources.txt - static cubin/PTX resource report.",
        "orbit_full_horizon_oracle_console.txt - stronger oracle failure evidence.",
        "compute_sanitizer_orbit_console.txt - zero-error memory-safety run.",
        "orbit_verification_metrics.csv - compact derived-metric summary.",
    ):
        add_list_item(doc, text, bullet_id)
    add_heading(doc, "Commands actually used", 2)
    add_table(
        doc,
        ("Purpose", "Command summary"),
        (
            ("Unmodified build", "powershell -ExecutionPolicy Bypass -File .\\scripts\\build_windows.ps1"),
            ("Pinned build", "cmake ... -T cuda=C:\\Program Files\\NVIDIA GPU Computing Toolkit\\CUDA\\v12.8 ..."),
            ("Seven-day demo", "demo_orbit_windows.ps1 -BuildDir build-cuda128-vs"),
            ("512 MiB", "stress_orbit_windows.ps1 -Preset laptop -BuildDir build-cuda128-vs"),
            ("2 GiB", "stress_orbit_windows.ps1 -Preset vram -BuildDir build-cuda128-vs"),
            ("Seed profile", "profile_orbit_windows.ps1 -BuildDir build-cuda128-vs"),
            ("Compact profile", "ncu --set full --kernel-name regex:query_seed_compact_kernel --launch-count 1 ... --mode all --write-events"),
            ("Full oracle", "klb_orbit_bench ... --preset file --verify-epochs 604800"),
        ),
        [1900, 7460],
    )
    add_heading(doc, "Overall assessment", 2)
    add_body(doc, "The v0.3 architecture turns the earlier launch-sized microbenchmark into a legitimate sustained crossover experiment. The direct seed path is real, correct under the stated acceptance gate, memory-efficient, narrowly faster than rebuilding a dense table for a one-off query, and demonstrably compute-bound. The design is promising as a query-first substrate, but it is not ready for operational orbital use until full-horizon numerical parity and SGP4 error/event-order validation are complete.")

    doc.save(OUTPUT)
    print(OUTPUT)
    print(SUMMARY_CSV)
    print(f"docx_sha256={hashlib.sha256(OUTPUT.read_bytes()).hexdigest()}")


if __name__ == "__main__":
    build_report()
