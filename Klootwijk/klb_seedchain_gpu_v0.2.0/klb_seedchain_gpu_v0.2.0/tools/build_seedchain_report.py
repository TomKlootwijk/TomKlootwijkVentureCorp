from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "KLB_SeedChain_GPU_v0.2.0_ELI5_Verification_Report.docx"

# standard_business_brief preset, plus named editorial-cover title overrides.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "1A1A1A"
MUTED = "5E6670"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
POSITIVE = "1F3A5F"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color, bold=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + side
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DBE2", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn("w:" + edge))
        if node is None:
            node = OxmlElement("w:" + edge)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    assert sum(widths_dxa) == CONTENT_DXA, (widths_dxa, sum(widths_dxa))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            cell.width = Inches(width / 1440.0)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.25, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    set_run_font(r, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths_dxa, numeric_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT_GRAY)
        align = WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=NAVY, size=9.1, align=align)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=9.05, align=align)
    set_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_numbering_definition(doc, *, bullet):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    # OOXML requires every abstractNum definition to precede every concrete
    # num instance. Appending an abstractNum after python-docx's built-in num
    # elements makes Word repair/ignore the custom definition, which caused
    # intended bullets to render as one continuously numbered list.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, *, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, *, bold_prefix=None, italic=False, color=INK, after=6, align=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)
    return p


def add_callout(doc, label, text, *, fill=CALLOUT, accent=POSITIVE):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [CONTENT_DXA])
    set_cell_shading(table.cell(0, 0), fill)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.10
    rl = p.add_run(label + "  ")
    set_run_font(rl, size=11, color=accent, bold=True)
    rt = p.add_run(text)
    set_run_font(rt, size=11, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_field(paragraph, instruction):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    rpr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    rpr.append(sz)
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, "Calibri", 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, "Calibri", 16, BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = False

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, "Calibri", 13, BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, "Calibri", 12, DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("Page ")
    set_run_font(r, size=9, color=MUTED)
    add_field(fp, "PAGE")
    r = fp.add_run(" of ")
    set_run_font(r, size=9, color=MUTED)
    add_field(fp, "NUMPAGES")

    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ffp.add_run("Independent test on RTX 5070 Ti Laptop GPU  |  16 August 2026")
    set_run_font(r, size=9, color=MUTED)

    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def start_body_section(doc):
    """Start a linked report section on a fresh page."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    return section


def add_running_header(doc):
    """Add a stable in-flow running label at an intentional page start."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("KLB SeedChain GPU 0.2.0  |  Independent verification")
    set_run_font(r, size=9, color=MUTED, bold=True)
    return p


def add_page_guard(doc, *, after=28):
    """Reserve top-edge space before page starts affected by Word parity layout."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(" ")
    set_run_font(r, size=1, color=WHITE)
    return p


def build_report():
    doc = Document()
    configure_document(doc)
    bullet_id = add_numbering_definition(doc, bullet=True)
    number_id = add_numbering_definition(doc, bullet=False)

    props = doc.core_properties
    props.title = "KLB SeedChain GPU 0.2.0 - ELI5 Verification Report"
    props.subject = "Independent CUDA, compression, correctness, and feasibility findings"
    props.author = "OpenAI Codex"
    props.keywords = "KLB, SeedChain, CUDA, compression, point cloud, verification"

    # Cover / editorial title block.
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(18)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    r = kicker.add_run("INDEPENDENT VERIFICATION REPORT")
    set_run_font(r, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("KLB SeedChain GPU 0.2.0")
    set_run_font(r, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("What works, what does not, and where it is practically useful")
    set_run_font(r, size=14, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(24)
    r = meta.add_run("Human-readable / ELI5 edition  |  RTX 5070 Ti Laptop  |  16 August 2026")
    set_run_font(r, size=10, color=MUTED, italic=True)

    add_callout(
        doc,
        "BOTTOM LINE",
        "The 207x demo compression is real for the supplied reconstructible sequence. The GPU implementation is correct and very fast. The same ratio does not carry over to arbitrary motion: the generic fitter became larger than dense float3 on a challenging 16-frame deformation test.",
        fill=LIGHT_BLUE,
        accent=POSITIVE,
    )

    metrics = doc.add_table(rows=2, cols=3)
    set_repeat_table_header(metrics.rows[0])
    set_table_geometry(metrics, [3120, 3120, 3120])
    labels = ("DEMO STORAGE", "FULL GPU CHECK", "DIRECT QUERY")
    values = ("207.321x", "0 lineage mismatches", "0.041 ms")
    for i in range(3):
        set_cell_shading(metrics.cell(0, i), LIGHT_GRAY)
        set_cell_text(metrics.cell(0, i), labels[i], bold=True, color=NAVY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(metrics.cell(1, i), values[i], bold=True, color=POSITIVE, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_body(
        doc,
        "ELI5 analogy: keep one master LEGO model, a tiny recipe card for each animation frame, and a short bag of replacement bricks only when something unpredictable changes.",
        italic=True,
        color=MUTED,
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    start_body_section(doc)
    add_running_header(doc)
    add_heading(doc, "1. The answer in plain English", 1)
    add_callout(
        doc,
        "VERDICT",
        "PASS for generated or strongly predictable point sequences. CONDITIONAL for real captured animation. FAIL as a claim of universal 200x geometry compression.",
        fill=CALLOUT,
        accent=POSITIVE,
    )
    add_body(
        doc,
        "The seed-chain idea works when most frames can be recreated from a known recipe. Instead of storing every point at every time step, it stores one compact starting shape, 96 bytes of instructions per frame, and sparse 16-byte correction records. That is why the supplied 240-frame chain shrinks from 188,743,680 bytes of dense float3 data to 910,392 bytes.",
    )
    add_body(
        doc,
        "The same trick is not magic compression. If almost every point needs a correction every frame, the correction log becomes larger than the dense sequence. That happened in the generic fitting challenge: 93.16% of point-frames needed novelty records, making the accurate chain 27% larger than dense float3.",
    )

    add_heading(doc, "What the implementation genuinely solves", 2)
    for text in (
        "Stores a long, reconstructible point sequence in a tiny container.",
        "Rebuilds any selected frame directly on an RTX 5070 Ti Laptop GPU.",
        "Queries compressed points without first materializing the complete dense sequence.",
        "Produces the same compact point/lineage event set as its materialized-frame path.",
        "Keeps unpredictable changes as sparse, checkpointed novelty records.",
    ):
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "What it does not solve", 2)
    for text in (
        "It is not a universal 200x compressor for arbitrary animation.",
        "It stores point positions only: no faces, normals, UVs, materials, or changing topology.",
        "It does not prove physical Klein-bottle memory, zero heat, or zero bandwidth.",
        "Its FNV integrity chain detects accidental change but is not a cryptographic signature.",
    ):
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "2. Test environment and scope", 1)
    add_table(
        doc,
        ("Item", "Measured configuration"),
        (
            ("GPU", "NVIDIA GeForce RTX 5070 Ti Laptop GPU, compute capability 12.0"),
            ("Driver / runtime", "NVIDIA driver 591.59; CUDA driver API 13.1; runtime 12.8"),
            ("Compiler", "CUDA 12.8.61, MSVC 19.44, native sm_120 plus compute_120 PTX"),
            ("Device", "46 SMs, 36 MiB reported L2, 192-bit memory bus, 11.94 GiB VRAM"),
            ("Power context", "Windows Balanced plan, AC power; no fixed vendor performance profile recorded"),
        ),
        [2200, 7160],
    )

    # The environment table naturally fills page 2; a second forced break here
    # creates a blank page in Word. Let normal pagination start section 3.
    add_page_guard(doc)
    add_running_header(doc)
    add_heading(doc, "3. How SeedChain works (ELI5)", 1)
    add_body(doc, "Imagine a flip-book containing 240 pictures of the same object:")
    steps = (
        "Store one compact master model. Each point uses a 37-bit log-spherical record.",
        "Store a recipe card for each picture. A 96-byte node contains the seed, motion predictor, grammar state, checkpoint information, and hash links.",
        "Store only surprises. A 16-byte novelty record is written when a point cannot be predicted closely enough.",
        "Jump to any frame. The GPU starts from the master, applies that frame's recipe, then walks back no farther than the nearest checkpoint to collect corrections.",
        "Ask a local question. The query tests cone/radius support, route compatibility, and an SDF guard, then appends only verified events.",
    )
    for text in steps:
        add_list_item(doc, text, number_id)

    add_callout(
        doc,
        "WHY 207x IS POSSIBLE",
        "A dense sequence repeats 65,536 coordinates 240 times. SeedChain stores the base once and explains almost all later motion procedurally. Only 0.232054% of point-frames need novelty records.",
        fill=LIGHT_BLUE,
        accent=POSITIVE,
    )

    add_heading(doc, "Storage equation", 2)
    add_body(doc, "KLSC1 bytes = 256-byte header + 96 x frames + 16 x novelty records + embedded base words.")
    add_body(doc, "Dense float3 bytes = 12 x points x frames.")
    add_body(doc, "For large sequences, the approximate break-even novelty density is below 75%. Useful compression requires much lower density: about 37.5% for 2x and 7.5% for 10x, before header/base overhead.")

    add_heading(doc, "4. Compression evidence", 1)
    add_table(
        doc,
        ("Case", "Frames", "Container", "vs float3", "Novelty"),
        (
            ("Included generated", "240", "910,392 B", "207.321x", "0.2321%"),
            ("Bunny + generated motion", "240", "508,856 B", "203.451x", "0.2308%"),
            ("Bunny generic fit, 0.2% max error", "16", "8,741,384 B", "0.7896x", "93.1562%"),
            ("Bunny generic fit, 4.0% max error", "16", "4,061,768 B", "1.6992x", "42.3043%"),
        ),
        [3060, 850, 1900, 1450, 2100],
        numeric_cols=(1, 2, 3, 4),
    )
    add_body(doc, "The Bunny tests preserve vertex positions only. The 240-frame Bunny chain still uses generated motion; it is not a fitted real scan sequence.", italic=True, color=MUTED)

    # Section 4 already fills page 3, so normal pagination is sufficient.
    add_page_guard(doc)
    add_running_header(doc)
    add_heading(doc, "5. The generic fitter reality check", 1)
    add_body(
        doc,
        "Sixteen frames were exported from the Bunny-based generated chain and then treated as an unknown external sequence. The fitter was allowed only one global uniform scale, one Y-axis rotation, and translation per frame. This deliberately tests whether the current generic predictor can explain complex per-point motion without knowing the original grammar.",
    )
    add_table(
        doc,
        ("Residual cutoff", "Novelty density", "Compression", "RMS/radius", "Max/radius"),
        (
            ("0.2%", "93.1562%", "0.7896x", "0.0334%", "0.2000%"),
            ("0.4%", "92.8466%", "0.7921x", "0.0516%", "0.4000%"),
            ("1.0%", "89.6210%", "0.8201x", "0.2475%", "1.0000%"),
            ("2.0%", "77.0262%", "0.9511x", "0.7075%", "2.0000%"),
            ("4.0%", "42.3043%", "1.6992x", "1.7981%", "4.0000%"),
        ),
        [1900, 1960, 1700, 1900, 1900],
        numeric_cols=(0, 1, 2, 3, 4),
    )
    add_callout(
        doc,
        "INTERPRETATION",
        "The fitter is accurate, but its predictor is too simple for this deformation. At the useful 0.2% maximum-error setting, storing corrections is more expensive than storing float3. Compression appears only after accepting a much larger 4% maximum positional error.",
        fill="FFF8E8",
        accent=GOLD,
    )
    add_body(
        doc,
        "The independent verify-sequence command confirmed 0.033386% RMS and 0.199999% maximum error for the accurate fit. This proves position reconstruction against the source PLYs. It does not yet prove that event membership versus the original source is unchanged.",
    )

    add_heading(doc, "6. GPU performance", 1)
    add_table(
        doc,
        ("Case", "Depth", "Seed query", "Dense query", "Penalty", "Events"),
        (
            ("Included, prescribed run", "15", "0.04124 ms", "0.01325 ms", "3.11x", "79"),
            ("Included, full-point verify", "15", "0.04101 ms", "0.01230 ms", "3.33x", "79"),
            ("Checkpoint frame 16", "0", "0.01343 ms", "0.01181 ms", "1.14x", "55"),
            ("Bunny generated", "15", "0.02672 ms", "0.01211 ms", "2.21x", "2,800"),
            ("Bunny fitted, accurate", "15", "0.05128 ms", "0.01249 ms", "4.10x", "2,761"),
            ("100% event yield", "15", "0.04105 ms", "0.01309 ms", "3.14x", "65,536"),
        ),
        [2480, 760, 1650, 1650, 1300, 1520],
        numeric_cols=(1, 2, 3, 4, 5),
    )

    start_body_section(doc)
    add_running_header(doc)
    add_heading(doc, "Repeatability at final frame (30 fresh processes)", 2)
    add_table(
        doc,
        ("Mode", "p50", "p95", "p99", "Mean"),
        (
            ("Decode one frame", "0.039165 ms", "0.039762 ms", "0.039858 ms", "0.039226 ms"),
            ("Direct seed query", "0.041012 ms", "0.043019 ms", "0.043067 ms", "0.041311 ms"),
            ("Dense query", "0.011993 ms", "0.013468 ms", "0.015226 ms", "0.012259 ms"),
        ),
        [2760, 1650, 1650, 1650, 1650],
        numeric_cols=(1, 2, 3, 4),
    )

    add_heading(doc, "7. What the timings mean", 1)
    add_body(
        doc,
        "The direct compressed path is slower than querying a frame that is already materialized. That is expected: it decodes 37-bit records, runs grammar/topology math, and searches up to 16 novelty ranges before evaluating the same event test.",
    )
    add_callout(
        doc,
        "ONE QUERY VS MANY",
        "At frame 239, one direct query costs about 0.041 ms. Decode-then-query costs about 0.039 + 0.012 = 0.051 ms, so direct wins for a one-off question. By roughly the second query against the same frame, materializing once and reusing the dense frame becomes faster, if the extra memory is acceptable.",
        fill=LIGHT_BLUE,
        accent=POSITIVE,
    )
    add_body(
        doc,
        "Checkpoint depth matters. Direct query time rose from roughly 0.013-0.015 ms at depth 0-1 to 0.039-0.043 ms at depth 15. This is the cost of dependent novelty lookups. Smaller checkpoint strides trade more file storage for faster random access.",
    )
    add_body(
        doc,
        "A 200,000-launch seed-query stress run sustained 0.041111 ms and 1.594 billion candidate points/s. After ramp-up, the GPU reported 98-99% utilization, about 128.6-130.6 W, and 63-69 C. Coarse nvidia-smi memory utilization rounded to 0%, consistent with this sub-1 MiB chain fitting in cache, but it is not a DRAM counter.",
    )

    start_body_section(doc)
    add_page_guard(doc, after=56)
    add_running_header(doc)
    add_heading(doc, "8. Correctness evidence", 1)
    checks = (
        "Package manifest: every listed file matched SHA-256; supplied KLSC chain hash matched the manifest.",
        "CPU oracle suite: 1 of 1 test executable passed, covering bit packing, parity, swizzle, Klein seam rules, PLY I/O, checkpoints, hashes, fitting, and error reporting.",
        "Full GPU comparison: all 65,536 points checked at frame 239; zero lineage mismatches and numerical error below the benchmark tolerance.",
        "Event equivalence: all 79 default events matched point identity and lineage; SDF and guard differences printed as 0.000000.",
        "Route partition: route 0 produced 38 events and route 1 produced 41, exactly partitioning the 79-event unfiltered set.",
        "High-yield path: 65,536 of 65,536 points appended; every compressed event matched the dense event set.",
        "Compute Sanitizer: zero memory-access errors across decode, direct query, dense query, and verification.",
        "CUDA compiler: native sm_120 kernels built with zero register spills; seed/decode uses 40 registers and dense query uses 18.",
    )
    for item in checks:
        add_list_item(doc, item, bullet_id)

    add_body(
        doc,
        "Important precision note: the benchmark prints CPU/GPU error with only three decimals. A displayed 0.000 means 'below the printed precision and acceptance threshold,' not necessarily bit-for-bit identical coordinates.",
        italic=True,
        color=MUTED,
    )

    start_body_section(doc)
    add_running_header(doc)
    add_heading(doc, "9. Practical use cases", 1)
    add_callout(
        doc,
        "BEST MATCH",
        "Use SeedChain where motion is mostly known and deterministic, edits are sparse, stable point identity matters, and you need occasional local GPU queries without keeping every dense frame resident.",
        fill=LIGHT_BLUE,
        accent=POSITIVE,
    )

    use_cases = (
        ("Procedural vegetation, particles, and branching effects", "Strong fit. Wind, growth, branching, and repeated cycles can be regenerated; artist edits become novelty records. Add cluster-level predictors for large assets."),
        ("Rollback/replay and deterministic simulation logs", "Strong fit when every participant consumes the same chain file. Store closed dynamics as seeds and only retain external interventions."),
        ("Sparse spatial event or culling service", "Strong fit for one-off frame queries that return few candidates: collision candidates, proximity alarms, local support checks, or render culling."),
        ("Repeated mechanical motion and digital twins", "Promising for rigid assemblies, conveyors, robot cells, and cyclic machines. Per-part rigid predictors are the obvious next extension."),
        ("Stable-correspondence LiDAR/depth sequences", "Conditional. Useful when scans are registered and most motion is rigid or cluster-rigid. Poor fit when correspondence changes, occlusion is heavy, or points are resampled."),
        ("Network or disk streaming to memory-limited GPUs", "Promising when sending one base plus compact nodes is cheaper than streaming dense frames. Chain segments can be prefetched and queried directly."),
    )
    add_heading(doc, "Suitable workloads", 2)
    for title_text, detail in use_cases:
        add_heading(doc, title_text, 3)
        add_body(doc, detail)

    add_heading(doc, "Poor matches", 2)
    for text in (
        "Cloth, fluids, crowds, or arbitrary deformation with the current single global predictor.",
        "Topology-changing meshes, unstable vertex correspondence, remeshing, or particle birth/death without an identity layer.",
        "Full production meshes requiring faces, normals, UVs, materials, skinning, or blendshape metadata.",
        "Workloads that query the same frame many times and can afford to materialize it once.",
        "Security/provenance systems that require cryptographic authentication rather than FNV integrity checks.",
    ):
        add_list_item(doc, text, bullet_id)

    # The use-case page is naturally full; avoid a redundant forced break.
    add_page_guard(doc)
    add_running_header(doc)
    add_heading(doc, "10. Limitations and issues found", 1)
    issues = (
        ("Universal compression is unproven", "The measured 207x applies to a reconstructible generated sequence. The generic fitting challenge was 0.79x at the accurate setting."),
        ("Dense-reference comparison is decode consistency", "The GPU benchmark compares direct reconstruction with a dense frame materialized from the same compressed chain. It does not compare event membership against the original source PLY sequence."),
        ("Hardware counters unavailable", "Nsight Compute was installed, but Windows denied performance-counter access with ERR_NVGPUCTRPERM. DRAM bytes, L2 hit rate, occupancy, stalls, and atomic serialization therefore remain unmeasured."),
        ("Cross-compiler regeneration is not byte-identical", "The supplied GCC-built chain and an MSVC regeneration had identical size/statistics but different terminal/SHA hashes. Two MSVC regenerations matched each other. Share the built chain or adopt canonical math for cross-platform deterministic generation."),
        ("Windows helper can report false success", "The supplied PowerShell build script continued after CMake failed on a stale CUDA 12.9 Visual Studio integration and ultimately exited 0. Pinning CUDA 12.8 in CMake produced a valid native build."),
        ("One-bit route is not identity", "Durable identity still requires point/base address, ordered node history, schema version, and novelty records."),
    )
    for label, detail in issues:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(label + ": ")
        set_run_font(r, bold=True, color=RED if "false" in label.lower() else NAVY)
        r = p.add_run(detail)
        set_run_font(r)

    start_body_section(doc)
    add_running_header(doc)
    add_heading(doc, "11. Recommended next engineering steps", 1)
    recommendations = (
        "Add a source-frame event oracle so fitted compression is accepted only when event membership/order also matches the original sequence.",
        "Replace one global Y-similarity predictor with per-cluster rigid/quaternion predictors and per-cluster quantization bounds.",
        "Sweep checkpoint strides 4, 8, 16, 32, and 64 on each real workload; jointly report file size, p95 latency, and novelty density.",
        "Enable NVIDIA performance counters and collect real DRAM/L2, occupancy, warp-stall, instruction, and atomic metrics.",
        "Make procedural generation cross-platform deterministic using canonical/fixed approximations or by distributing authoritative node data.",
        "Fix the Windows build helper to stop on non-zero native exit codes and optionally select an installed CUDA toolkit explicitly.",
        "Compare against conventional point-cloud, geometry, and video codecs at equal error and required random-access behavior.",
        "Add stable face/index streams only for applications that truly need mesh output; keep the direct event path position-only when that is sufficient.",
    )
    number_id = add_numbering_definition(doc, bullet=False)
    for item in recommendations:
        add_list_item(doc, item, number_id)

    add_callout(
        doc,
        "GO / NO-GO RULE",
        "Proceed when novelty stays sparse, maximum error is below the application guard margin, source-event classifications are preserved, and avoided storage/transfer is worth the 1.1x-4.1x direct-query compute penalty. Otherwise materialize, segment, or use a conventional codec.",
        fill="FFF8E8",
        accent=GOLD,
    )

    start_body_section(doc)
    add_page_guard(doc, after=76)
    add_running_header(doc)
    add_heading(doc, "Appendix A. Evidence and commands", 1)
    add_body(doc, "Primary measured artifacts created in the project directory:")
    for text in (
        "seedchain_results.csv - prescribed 20-repeat final-frame run.",
        "seedchain_depth_sweep.csv - checkpoint-depth sweep across nine frames.",
        "seedchain_process_runs.csv - 30 fresh-process stability runs.",
        "seedchain_full_verify_results.csv - all 65,536 points verified.",
        "bunny_seedchain_results.csv - real Bunny geometry with generated motion.",
        "fitted_sequence_results.csv - accurate and loose generic-fit GPU runs.",
        "seedchain_high_yield.csv - 100% event-yield compaction test.",
    ):
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "Core commands", 2)
    command_rows = (
        ("Build", 'cmake -S . -B build-cuda128 -G "Visual Studio 17 2022" -A x64 -T "cuda=C:\\Program Files\\NVIDIA GPU Computing Toolkit\\CUDA\\v12.8" -DKLB_CUDA_ARCH=120 -DKLB_REQUIRE_CUDA=ON'),
        ("CPU tests", "ctest --test-dir build-cuda128 -C Release --output-on-failure"),
        ("Inspect", "klb_seedchain.exe inspect data\\procedural_65536_240f.klsc"),
        ("GPU verify", "klb_seedchain_bench.exe data\\procedural_65536_240f.klsc --frame 239 --mode all --repeats 200 --verify 65536"),
        ("Sanitizer", "compute-sanitizer --tool memcheck klb_seedchain_bench.exe ..."),
        ("Fitter verify", "klb_seedchain.exe verify-sequence bunny_frames_16.txt bunny_fitted_16f_thr002.klsc"),
    )
    add_table(doc, ("Purpose", "Command summary"), command_rows, [1700, 7660])

    add_heading(doc, "Evidence boundary", 2)
    add_body(
        doc,
        "The report is based on the supplied package, locally generated test artifacts, and direct execution on the stated laptop. The Stanford Bunny adapter uses vertices only. No claim is made about faces, materials, visual perceptual quality, adversarial integrity, or source-level event preservation beyond the tests described.",
    )
    add_body(
        doc,
        "Overall assessment: the architecture is a credible specialized seed-plus-novelty event substrate. Its value comes from matching the right predictable workload, not from treating the demo compression ratio as universal.",
        bold_prefix="Overall assessment:",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
